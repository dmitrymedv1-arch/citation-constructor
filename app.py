import os
import csv
import streamlit as st
import re
import json
from datetime import datetime
from crossref.restful import Works
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from tqdm import tqdm
from docx.oxml import OxmlElement
import base64
import html
import concurrent.futures
from typing import List, Dict, Tuple, Set, Any, Optional
import hashlib
import time
from collections import Counter
import functools
import logging
from pathlib import Path
import sqlite3
from contextlib import contextmanager
import requests

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('citation_processor.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Конфигурация
class Config:
    """Конфигурационные константы приложения"""
    # Пути к файлам
    DB_PATH = "doi_cache.db"
    LTWA_CSV_PATH = "ltwa.csv"
    USER_PREFS_DB = "user_preferences.db"
    
    # Настройки API
    CROSSREF_WORKERS = 3
    CROSSREF_RETRY_WORKERS = 2
    REQUEST_TIMEOUT = 30
    
    # Кэширование
    CACHE_TTL_HOURS = 24 * 7  # 1 неделя
    
    # Валидация
    MIN_REFERENCES_FOR_STATS = 5
    MAX_REFERENCES = 1000
    
    # Стили
    NUMBERING_STYLES = ["No numbering", "1", "1.", "1)", "(1)", "[1]"]
    AUTHOR_FORMATS = ["AA Smith", "A.A. Smith", "Smith AA", "Smith A.A", "Smith, A.A."]
    PAGE_FORMATS = ["122 - 128", "122-128", "122 – 128", "122–128", "122–8", "122"]
    DOI_FORMATS = ["10.10/xxx", "doi:10.10/xxx", "DOI:10.10/xxx", "https://dx.doi.org/10.10/xxx"]
    JOURNAL_STYLES = ["{Full Journal Name}", "{J. Abbr.}", "{J Abbr}"]
    AVAILABLE_ELEMENTS = ["", "Authors", "Title", "Journal", "Year", "Volume", "Issue", "Pages", "DOI"]
    
    # Цвета прогресс-бара
    PROGRESS_COLORS = {
        'start': '#FF6B6B',
        'middle': '#4ECDC4', 
        'end': '#45B7D1'
    }
    
    # Настройки тем (обновленные для лучшего контраста)
    THEMES = {
        'light': {
            'primary': '#1f77b4',
            'background': '#f8f9fa',
            'secondaryBackground': '#ffffff',
            'text': '#212529',
            'font': 'sans-serif',
            'border': '#dee2e6',
            'cardBackground': '#ffffff'
        },
        'dark': {
            'primary': '#4ECDC4',
            'background': '#1a1d23',
            'secondaryBackground': '#2d323d',
            'text': '#e9ecef',
            'font': 'sans-serif',
            'border': '#495057',
            'cardBackground': '#2d323d'
        }
    }

# Инициализация Crossref
works = Works()

# Полный словарь переводов
TRANSLATIONS = {
    'en': {
        'header': '🎨 Citation Style Constructor',
        'general_settings': '⚙️ General Settings',
        'element_config': '📑 Element Configuration',
        'style_preview': '👀 Style Preview',
        'data_input': '📁 Data Input',
        'data_output': '📤 Data Output',
        'numbering_style': 'Numbering:',
        'author_format': 'Authors:',
        'author_separator': 'Separator:',
        'et_al_limit': 'Et al after:',
        'use_and': "'and'",
        'use_ampersand': "'&'",
        'doi_format': 'DOI format:',
        'doi_hyperlink': 'DOI as hyperlink',
        'page_format': 'Pages:',
        'final_punctuation': 'Final punctuation:',
        'element': 'Element',
        'italic': 'Italic',
        'bold': 'Bold',
        'parentheses': 'Parentheses',
        'separator': 'Separator',
        'input_method': 'Input:',
        'output_method': 'Output:',
        'select_docx': 'Select DOCX',
        'enter_references': 'Enter references (one per line)',
        'references': 'References:',
        'results': 'Results:',
        'process': '🚀 Process',
        'example': 'Example:',
        'error_select_element': 'Select at least one element!',
        'processing': '⏳ Processing...',
        'upload_file': 'Upload a file!',
        'enter_references_error': 'Enter references!',
        'select_docx_output': 'Select DOCX output to download!',
        'doi_txt': '📄 DOI (TXT)',
        'references_docx': '📋 References (DOCX)',
        'found_references': 'Found {} references.',
        'found_references_text': 'Found {} references in text.',
        'statistics': 'Statistics: {} DOI found, {} not found.',
        'language': 'Language:',
        'gost_style': 'Apply GOST Style',
        'export_style': '📤 Export Style',
        'import_style': '📥 Import Style',
        'export_file_name': 'File name:',
        'import_file': 'Select style file:',
        'export_success': 'Style exported successfully!',
        'import_success': 'Style imported successfully!',
        'import_error': 'Error importing style file!',
        'processing_status': 'Processing references...',
        'current_reference': 'Current: {}',
        'processed_stats': 'Processed: {}/{} | Found: {} | Errors: {}',
        'time_remaining': 'Estimated time remaining: {}',
        'duplicate_reference': '🔄 Repeated Reference (See #{})',
        'batch_processing': 'Batch processing DOI...',
        'extracting_metadata': 'Extracting metadata...',
        'checking_duplicates': 'Checking for duplicates...',
        'retrying_failed': 'Retrying failed DOI requests...',
        'bibliographic_search': 'Searching by bibliographic data...',
        'style_presets': 'Style Presets',
        'gost_button': 'GOST',
        'acs_button': 'ACS (MDPI)',
        'rsc_button': 'RSC',
        'cta_button': 'CTA',
        'style_preset_tooltip': 'Here are some styles maintained by individual publishers. For major publishers (Elsevier, Springer Nature, and Wiley), styles vary from journal to journal. To create (or reformat) references for a specific journal, use the Citation Style Constructor.',
        'journal_style': 'Journal style:',
        'full_journal_name': 'Full Journal Name',
        'journal_abbr_with_dots': 'J. Abbr.',
        'journal_abbr_no_dots': 'J Abbr',
        'short_guide_title': 'A short guide for the conversion of doi-based references',
        'step_1': '❶ Select a ready reference style (ACS(MDPI), RSC, or CTA), or create your own style by selecting the sequence, design, and punctuation of the element configurations',
        'step_1_note': '(!) The punctuation boxes enable various items to be included between element configurations (simple punctuation, Vol., Issue…)',
        'step_2': '❷ Then, use the Style Presets to change certain element configurations for each reformatted reference.',
        'step_3': '❸ The Style Preview function enables users to visualize the final form of their reference style',
        'step_4': '❹ If the final style is appropriate, select the Docx or Text option in the Data Input section and upload the corresponding information (reference list). Then, in the Data Output section, select the required options and press "Process" to initiate reformatting.',
        'step_5': '❺ After processing is complete, download the reformatted references in your preferred format.',
        'step_5_note': '(!) Outputting the Docx file is recommended, as it preserves formatting (e.g., bold, italic, and hyperlinks) and includes additional stats at the end of the document.',
        'step_6': '❻ After creating your final version of the style, save it so that you can upload it again in the next session. Use the Style Management section for this purpose.',
        'validation_error_no_elements': 'Please configure at least one element or select a preset style!',
        'validation_error_too_many_references': 'Too many references (maximum {} allowed)',
        'validation_warning_few_references': 'Few references for meaningful statistics',
        'cache_initialized': 'Cache initialized successfully',
        'cache_cleared': 'Cache cleared successfully',
        'theme_selector': 'Theme:',
        'light_theme': 'Light',
        'dark_theme': 'Dark',
        'mobile_view': 'Mobile View',
        'desktop_view': 'Desktop View',
        'clear_button': '🗑️ Clear',
        'back_button': '↩️ Back'
    },
    'ru': {
        'header': '🎨 Конструктор стилей цитирования',
        'general_settings': '⚙️ Настройки',
        'element_config': '📑 Конфигурация элементов',
        'style_preview': '👀 Предпросмотр',
        'data_input': '📁 Ввод',
        'data_output': '📤 Вывод',
        'numbering_style': 'Нумерация:',
        'author_format': 'Авторы:',
        'author_separator': 'Разделитель:',
        'et_al_limit': 'Et al после:',
        'use_and': "'и'",
        'use_ampersand': "'&'",
        'doi_format': 'Формат DOI:',
        'doi_hyperlink': 'DOI как ссылка',
        'page_format': 'Страницы:',
        'final_punctuation': 'Конечная пунктуация:',
        'element': 'Элемент',
        'italic': 'Курсив',
        'bold': 'Жирный',
        'parentheses': 'Скобки',
        'separator': 'Разделитель',
        'input_method': 'Ввод:',
        'output_method': 'Вывод:',
        'select_docx': 'Выберите DOCX',
        'enter_references': 'Введите ссылки (по одной на строку)',
        'references': 'Ссылки:',
        'results': 'Результаты:',
        'process': '🚀 Обработать',
        'example': 'Пример:',
        'error_select_element': 'Выберите хотя бы один элемент!',
        'processing': '⏳ Обработка...',
        'upload_file': 'Загрузите файл!',
        'enter_references_error': 'Введите ссылки!',
        'select_docx_output': 'Выберите DOCX для скачивания!',
        'doi_txt': '📄 DOI (TXT)',
        'references_docx': '📋 Ссылки (DOCX)',
        'found_references': 'Найдено {} ссылок.',
        'found_references_text': 'Найдено {} ссылок в тексте.',
        'statistics': 'Статистика: {} DOI найдено, {} не найдено.',
        'language': 'Язык:',
        'gost_style': 'Применить стиль ГОСТ',
        'export_style': '📤 Экспорт стиля',
        'import_style': '📥 Импорт стиля',
        'export_file_name': 'Имя файла:',
        'import_file': 'Выберите файл стиля:',
        'export_success': 'Стиль экспортирован успешно!',
        'import_success': 'Стиль импортирован успешно!',
        'import_error': 'Ошибка импорта файла стиля!',
        'processing_status': 'Обработка ссылок...',
        'current_reference': 'Текущая: {}',
        'processed_stats': 'Обработано: {}/{} | Найдено: {} | Ошибки: {}',
        'time_remaining': 'Примерное время до завершения: {}',
        'duplicate_reference': '🔄 Повторная ссылка (См. #{})',
        'batch_processing': 'Пакетная обработка DOI...',
        'extracting_metadata': 'Извлечение метаданных...',
        'checking_duplicates': 'Проверка на дубликаты...',
        'retrying_failed': 'Повторная попытка для неудачных DOI...',
        'bibliographic_search': 'Поиск по библиографическим данным...',
        'style_presets': 'Готовые стили',
        'gost_button': 'ГОСТ',
        'acs_button': 'ACS (MDPI)',
        'rsc_button': 'RSC',
        'cta_button': 'CTA',
        'style_preset_tooltip': 'Здесь указаны некоторые стили, которые сохраняются в пределах одного издательства. Для ряда крупных издательств (Esevier, Springer Nature, Wiley) стиль отличается от журнала к журналу. Для формирования (или переформатирования) ссылок для конкретного журнала предлагаем воспользоваться конструктором ссылок.',
        'journal_style': 'Стиль журнала:',
        'full_journal_name': 'Полное название журнала',
        'journal_abbr_with_dots': 'J. Abbr.',
        'journal_abbr_no_dots': 'J Abbr',
        'short_guide_title': 'Краткое руководство для конвертации ссылок, имеющих doi',
        'step_1': '❶ Выберите готовый стиль ссылок (ГОСТ, ACS(MDPI), RSC или CTA) или создайте свой собственный стиль, выбрав последовательность, оформление и пунктуацию конфигураций элементов',
        'step_1_note': '(!) Поля пунктуации позволяют включать различные элементы между конфигурациями (простая пунктуация, Том, Выпуск…)',
        'step_2': '❷ Затем используйте готовые стили, чтобы изменить определенные конфигурации элементов для каждой переформатированной ссылки.',
        'step_3': '❸ Функция предпросмотра стиля позволяет визуализировать окончательную форму вашего стиля ссылок',
        'step_4': '❹ Если окончательный стиль подходит, выберите опцию Docx или Текст в разделе ввода данных и загрузите соответствующую информацию (список литературы). Затем в разделе вывода данных выберите нужные опции и нажмите "Обработать" для начала переформатирования.',
        'step_5': '❺ После завершения обработки загрузите переформатированные ссылки в предпочитаемом формате.',
        'step_5_note': '(!) Рекомендуется выводить файл Docx, так как он сохраняет форматирование (например, жирный шрифт, курсив и гиперссылки) и включает дополнительную статистику в конце документа.',
        'step_6': '❻ После создания окончательной версии стиля сохраните его, чтобы можно было снова загрузить в следующей сессии. Для этого используйте раздел Style Management.',
        'validation_error_no_elements': 'Пожалуйста, настройте хотя бы один элемент или выберите готовый стиль!',
        'validation_error_too_many_references': 'Слишком много ссылок (максимум {} разрешено)',
        'validation_warning_few_references': 'Мало ссылок для значимой статистики',
        'cache_initialized': 'Кэш инициализирован успешно',
        'cache_cleared': 'Кэш очищен успешно',
        'theme_selector': 'Тема:',
        'light_theme': 'Светлая',
        'dark_theme': 'Тёмная',
        'mobile_view': 'Мобильный вид',
        'desktop_view': 'Десктопный вид',
        'clear_button': '🗑️ Очистить',
        'back_button': '↩️ Назад'
    },
    'de': {
        'header': '🎨 Zitationsstil-Konstruktor',
        'general_settings': '⚙️ Allgemeine Einstellungen',
        'element_config': '📑 Elementkonfiguration',
        'style_preview': '👀 Stilvorschau',
        'data_input': '📁 Dateneingabe',
        'data_output': '📤 Datenausgabe',
        'numbering_style': 'Nummerierung:',
        'author_format': 'Autoren:',
        'author_separator': 'Trennzeichen:',
        'et_al_limit': 'Et al nach:',
        'use_and': "'und'",
        'use_ampersand': "'&'",
        'doi_format': 'DOI-Format:',
        'doi_hyperlink': 'DOI als Hyperlink',
        'page_format': 'Seiten:',
        'final_punctuation': 'Schlusszeichen:',
        'element': 'Element',
        'italic': 'Kursiv',
        'bold': 'Fett',
        'parentheses': 'Klammern',
        'separator': 'Trennzeichen',
        'input_method': 'Eingabe:',
        'output_method': 'Ausgabe:',
        'select_docx': 'DOCX auswählen',
        'enter_references': 'Referenzen eingeben (eine pro Zeile)',
        'references': 'Referenzen:',
        'results': 'Ergebnisse:',
        'process': '🚀 Verarbeiten',
        'example': 'Beispiel:',
        'error_select_element': 'Wählen Sie mindestens ein Element aus!',
        'processing': '⏳ Verarbeitung...',
        'upload_file': 'Laden Sie eine Datei hoch!',
        'enter_references_error': 'Geben Sie Referenzen ein!',
        'select_docx_output': 'Wählen Sie DOCX-Ausgabe zum Herunterladen!',
        'doi_txt': '📄 DOI (TXT)',
        'references_docx': '📋 Referenzen (DOCX)',
        'found_references': '{} Referenzen gefunden.',
        'found_references_text': '{} Referenzen im Text gefunden.',
        'statistics': 'Statistik: {} DOI gefunden, {} nicht gefunden.',
        'language': 'Sprache:',
        'gost_style': 'GOST-Stil anwenden',
        'export_style': '📤 Stil exportieren',
        'import_style': '📥 Stil importieren',
        'export_file_name': 'Dateiname:',
        'import_file': 'Stildatei auswählen:',
        'export_success': 'Stil erfolgreich exportiert!',
        'import_success': 'Stil erfolgreich importiert!',
        'import_error': 'Fehler beim Importieren der Stildatei!',
        'processing_status': 'Verarbeite Referenzen...',
        'current_reference': 'Aktuell: {}',
        'processed_stats': 'Verarbeitet: {}/{} | Gefunden: {} | Fehler: {}',
        'time_remaining': 'Geschätzte verbleibende Zeit: {}',
        'duplicate_reference': '🔄 Wiederholte Referenz (Siehe #{})',
        'batch_processing': 'Stapelverarbeitung DOI...',
        'extracting_metadata': 'Extrahiere Metadaten...',
        'checking_duplicates': 'Prüfe auf Duplikate...',
        'retrying_failed': 'Wiederhole fehlgeschlagene DOI-Anfragen...',
        'bibliographic_search': 'Suche nach bibliografischen Daten...',
        'style_presets': 'Stilvorlagen',
        'gost_button': 'GOST',
        'acs_button': 'ACS (MDPI)',
        'rsc_button': 'RSC',
        'cta_button': 'CTA',
        'style_preset_tooltip': 'Hier sind einige Stile, die von einzelnen Verlagen gepflegt werden. Für große Verlage (Elsevier, Springer Nature, Wiley) variiert der Stil von Journal zu Journal. Um Referenzen für ein bestimmtes Journal zu erstellen (oder neu zu formatieren), verwenden Sie den Zitationsstil-Konstruktor.',
        'journal_style': 'Journal-Stil:',
        'full_journal_name': 'Vollständiger Journalname',
        'journal_abbr_with_dots': 'J. Abk.',
        'journal_abbr_no_dots': 'J Abk',
        'short_guide_title': 'Kurzanleitung zur Konvertierung von DOI-basierten Referenzen',
        'step_1': '❶ Wählen Sie einen vorgefertigten Referenzstil (ACS(MDPI), RSC oder CTA) oder erstellen Sie Ihren eigenen Stil, indem Sie die Reihenfolge, Gestaltung und Zeichensetzung der Elementkonfigurationen auswählen',
        'step_1_note': '(!) Die Zeichensetzungsfelder ermöglichen die Aufnahme verschiedener Elemente zwischen Elementkonfigurationen (einfache Zeichensetzung, Vol., Issue…)',
        'step_2': '❷ Verwenden Sie dann die Stilvorlagen, um bestimmte Elementkonfigurationen für jede neu formatierte Referenz zu ändern.',
        'step_3': '❸ Die Stilvorschau-Funktion ermöglicht es Benutzern, die endgültige Form ihres Referenzstils zu visualisieren',
        'step_4': '❹ Wenn der endgültige Stil geeignet ist, wählen Sie die Option Docx oder Text im Abschnitt Dateneingabe und laden Sie die entsprechenden Informationen hoch. Wählen Sie dann im Abschnitt Datenausgabe die erforderlichen Optionen und drücken Sie "Verarbeiten", um die Neuformatierung zu starten.',
        'step_5': '❺ Nach Abschluss der Verarbeitung laden Sie die neu formatierten Referenzen in Ihrem bevorzugten Format herunter.',
        'step_5_note': '(!) Die Ausgabe der Docx-Datei wird empfohlen, da sie die Formatierung beibehält (z.B. fett, kursiv und Hyperlinks) und zusätzliche Statistiken am Ende des Dokuments enthält.',
        'step_6': '❻ Speichern Sie Ihre endgültige Version des Stils, damit Sie ihn in der nächsten Sitzung erneut hochladen können. Verwenden Sie dazu den Abschnitt Stilmanagement.',
        'validation_error_no_elements': 'Bitte konfigurieren Sie mindestens ein Element oder wählen Sie einen vorgefertigten Stil!',
        'validation_error_too_many_references': 'Zu viele Referenzen (maximal {} erlaubt)',
        'validation_warning_few_references': 'Wenige Referenzen für aussagekräftige Statistiken',
        'cache_initialized': 'Cache erfolgreich initialisiert',
        'cache_cleared': 'Cache erfolgreich gelöscht',
        'theme_selector': 'Thema:',
        'light_theme': 'Hell',
        'dark_theme': 'Dunkel',
        'mobile_view': 'Mobile Ansicht',
        'desktop_view': 'Desktop Ansicht',
        'clear_button': '🗑️ Löschen',
        'back_button': '↩️ Zurück'
    },
    'es': {
        'header': '🎨 Constructor de Estilos de Citas',
        'general_settings': '⚙️ Configuración General',
        'element_config': '📑 Configuración de Elementos',
        'style_preview': '👀 Vista Previa del Estilo',
        'data_input': '📁 Entrada de Datos',
        'data_output': '📤 Salida de Datos',
        'numbering_style': 'Numeración:',
        'author_format': 'Autores:',
        'author_separator': 'Separador:',
        'et_al_limit': 'Et al después de:',
        'use_and': "'y'",
        'use_ampersand': "'&'",
        'doi_format': 'Formato DOI:',
        'doi_hyperlink': 'DOI como hipervínculo',
        'page_format': 'Páginas:',
        'final_punctuation': 'Puntuación final:',
        'element': 'Elemento',
        'italic': 'Cursiva',
        'bold': 'Negrita',
        'parentheses': 'Paréntesis',
        'separator': 'Separador',
        'input_method': 'Entrada:',
        'output_method': 'Salida:',
        'select_docx': 'Seleccionar DOCX',
        'enter_references': 'Ingresar referencias (una por línea)',
        'references': 'Referencias:',
        'results': 'Resultados:',
        'process': '🚀 Procesar',
        'example': 'Ejemplo:',
        'error_select_element': '¡Seleccione al menos un elemento!',
        'processing': '⏳ Procesando...',
        'upload_file': '¡Suba un archivo!',
        'enter_references_error': '¡Ingrese referencias!',
        'select_docx_output': '¡Seleccione salida DOCX para descargar!',
        'doi_txt': '📄 DOI (TXT)',
        'references_docx': '📋 Referencias (DOCX)',
        'found_references': 'Se encontraron {} referencias.',
        'found_references_text': 'Se encontraron {} referencias en el texto.',
        'statistics': 'Estadísticas: {} DOI encontrados, {} no encontrados.',
        'language': 'Idioma:',
        'gost_style': 'Aplicar Estilo GOST',
        'export_style': '📤 Exportar Estilo',
        'import_style': '📥 Importar Estilo',
        'export_file_name': 'Nombre del archivo:',
        'import_file': 'Seleccionar archivo de estilo:',
        'export_success': '¡Estilo exportado exitosamente!',
        'import_success': '¡Estilo importado exitosamente!',
        'import_error': '¡Error al importar archivo de estilo!',
        'processing_status': 'Procesando referencias...',
        'current_reference': 'Actual: {}',
        'processed_stats': 'Procesadas: {}/{} | Encontradas: {} | Errores: {}',
        'time_remaining': 'Tiempo restante estimado: {}',
        'duplicate_reference': '🔄 Referencia Repetida (Ver #{})',
        'batch_processing': 'Procesamiento por lotes DOI...',
        'extracting_metadata': 'Extrayendo metadatos...',
        'checking_duplicates': 'Verificando duplicados...',
        'retrying_failed': 'Reintentando solicitudes DOI fallidas...',
        'bibliographic_search': 'Buscando por datos bibliográficos...',
        'style_presets': 'Estilos Predefinidos',
        'gost_button': 'GOST',
        'acs_button': 'ACS (MDPI)',
        'rsc_button': 'RSC',
        'cta_button': 'CTA',
        'style_preset_tooltip': 'Aquí hay algunos estilos mantenidos por editoriales individuales. Para editoriales importantes (Elsevier, Springer Nature, Wiley), el estilo varía de revista en revista. Para crear (or reformatear) referencias para una revista específica, use el Constructor de Estilos de Citas.',
        'journal_style': 'Estilo de revista:',
        'full_journal_name': 'Nombre Completo de la Revista',
        'journal_abbr_with_dots': 'J. Abrev.',
        'journal_abbr_no_dots': 'J Abrev',
        'short_guide_title': 'Una guía breve para la conversión de referencias basadas en doi',
        'step_1': '❶ Seleccione un estilo de referencia listo (ACS(MDPI), RSC o CTA), o cree su propio estilo seleccionando la secuencia, diseño y puntuación de las configuraciones de elementos',
        'step_1_note': '(!) Los cuadros de puntuación permiten incluir varios elementos entre configuraciones de elementos (puntuación simple, Vol., Issue…)',
        'step_2': '❷ Luego, use los Estilos Predefinidos para cambiar ciertas configuraciones de elementos para cada referencia reformateada.',
        'step_3': '❸ La función de Vista Previa del Estilo permite a los usuarios visualizar la forma final de su estilo de referencia',
        'step_4': '❹ Si el estilo final es apropiado, seleccione la opción Docx o Texto en la sección de Entrada de Datos y cargue la información correspondiente. Luego, en la sección de Salida de Datos, seleccione las opciones requeridas y presione "Procesar" para iniciar el reformateo.',
        'step_5': '❺ Después de completar el procesamiento, descargue las referencias reformateadas en su formato preferido.',
        'step_5_note': '(!) Se recomienda generar el archivo Docx, ya que conserva el formato (por ejemplo, negrita, cursiva e hipervínculos) e incluye estadísticas adicionales al final del documento.',
        'step_6': '❻ Después de crear su versión final del estilo, guárdela para poder cargarla nuevamente en la siguiente sesión. Use la sección Gestión de Estilos para este propósito.',
        'validation_error_no_elements': '¡Por favor configure al menos un elemento o seleccione un estilo predefinido!',
        'validation_error_too_many_references': 'Demasiadas referencias (máximo {} permitidas)',
        'validation_warning_few_references': 'Pocas referencias para estadísticas significativas',
        'cache_initialized': 'Caché inicializado exitosamente',
        'cache_cleared': 'Caché limpiado exitosamente',
        'theme_selector': 'Tema:',
        'light_theme': 'Claro',
        'dark_theme': 'Oscuro',
        'mobile_view': 'Vista Móvil',
        'desktop_view': 'Vista Escritorio',
        'clear_button': '🗑️ Limpiar',
        'back_button': '↩️ Atrás'
    },
    'it': {
        'header': '🎨 Costruttore di Stili di Citazione',
        'general_settings': '⚙️ Impostazioni Generali',
        'element_config': '📑 Configurazione Elementi',
        'style_preview': '👀 Anteprima Stile',
        'data_input': '📁 Input Dati',
        'data_output': '📤 Output Dati',
        'numbering_style': 'Numerazione:',
        'author_format': 'Autori:',
        'author_separator': 'Separatore:',
        'et_al_limit': 'Et al dopo:',
        'use_and': "'e'",
        'use_ampersand': "'&'",
        'doi_format': 'Formato DOI:',
        'doi_hyperlink': 'DOI come collegamento ipertestuale',
        'page_format': 'Pagine:',
        'final_punctuation': 'Punteggiatura finale:',
        'element': 'Elemento',
        'italic': 'Corsivo',
        'bold': 'Grassetto',
        'parentheses': 'Parentesi',
        'separator': 'Separatore',
        'input_method': 'Input:',
        'output_method': 'Output:',
        'select_docx': 'Seleziona DOCX',
        'enter_references': 'Inserisci riferimenti (uno per riga)',
        'references': 'Riferimenti:',
        'results': 'Risultati:',
        'process': '🚀 Elabora',
        'example': 'Esempio:',
        'error_select_element': 'Seleziona almeno un elemento!',
        'processing': '⏳ Elaborazione...',
        'upload_file': 'Carica un file!',
        'enter_references_error': 'Inserisci i riferimenti!',
        'select_docx_output': 'Seleziona output DOCX da scaricare!',
        'doi_txt': '📄 DOI (TXT)',
        'references_docx': '📋 Riferimenti (DOCX)',
        'found_references': 'Trovati {} riferimenti.',
        'found_references_text': 'Trovati {} riferimenti nel testo.',
        'statistics': 'Statistiche: {} DOI trovati, {} non trovati.',
        'language': 'Lingua:',
        'gost_style': 'Applica Stile GOST',
        'export_style': '📤 Esporta Stile',
        'import_style': '📥 Importa Stile',
        'export_file_name': 'Nome file:',
        'import_file': 'Seleziona file stile:',
        'export_success': 'Stile esportato con successo!',
        'import_success': 'Stile importato con successo!',
        'import_error': 'Errore durante l\'importazione del file stile!',
        'processing_status': 'Elaborazione riferimenti...',
        'current_reference': 'Attuale: {}',
        'processed_stats': 'Elaborati: {}/{} | Trovati: {} | Errori: {}',
        'time_remaining': 'Tempo rimanente stimato: {}',
        'duplicate_reference': '🔄 Riferimento Ripetuto (Vedi #{})',
        'batch_processing': 'Elaborazione batch DOI...',
        'extracting_metadata': 'Estrazione metadati...',
        'checking_duplicates': 'Controllo duplicati...',
        'retrying_failed': 'Riprova richieste DOI fallite...',
        'bibliographic_search': 'Ricerca per dati bibliografici...',
        'style_presets': 'Stili Preimpostati',
        'gost_button': 'GOST',
        'acs_button': 'ACS (MDPI)',
        'rsc_button': 'RSC',
        'cta_button': 'CTA',
        'style_preset_tooltip': 'Ecco alcuni stili mantenuti da singoli editori. Per gli editori principali (Elsevier, Springer Nature, Wiley), lo stile varia da rivista a rivista. Per creare (o riformattare) riferimenti per una rivista specifica, utilizza il Costruttore di Stili di Citazione.',
        'journal_style': 'Stile rivista:',
        'full_journal_name': 'Nome Completo Rivista',
        'journal_abbr_with_dots': 'Riv. Abbr.',
        'journal_abbr_no_dots': 'Riv Abbr',
        'short_guide_title': 'Una breve guida per la conversione di riferimenti basati su doi',
        'step_1': '❶ Seleziona uno stile di riferimento pronto (ACS(MDPI), RSC o CTA), o crea il tuo stile personalizzato selezionando la sequenza, il design e la punteggiatura delle configurazioni degli elementi',
        'step_1_note': '(!) Le caselle di punteggiatura consentono di includere vari elementi tra le configurazioni degli elementi (punteggiatura semplice, Vol., Issue…)',
        'step_2': '❷ Quindi, utilizza gli Stili Preimpostati per modificare determinate configurazioni di elementi per ogni riferimento riformattato.',
        'step_3': '❸ La funzione Anteprima Stile consente agli utenti di visualizzare la forma finale del loro stile di riferimento',
        'step_4': '❹ Se lo stile finale è appropriato, seleziona l\'opzione Docx o Testo nella sezione Input Dati e carica le informazioni corrispondenti. Quindi, nella sezione Output Dati, seleziona le opzioni richieste e premi "Elabora" per avviare la riformattazione.',
        'step_5': '❺ Dopo il completamento dell\'elaborazione, scarica i riferimenti riformattati nel formato preferito.',
        'step_5_note': '(!) Si consiglia di output il file Docx, in quanto conserva la formattazione (ad esempio, grassetto, cursivo e collegamenti ipertestuali) e include statistiche aggiuntive alla fine del documento.',
        'step_6': '❻ Dopo aver creato la versione finale dello stile, salvala in modo da poterla caricare nuovamente nella sessione successiva. Utilizza la sezione Gestione Stili per questo scopo.',
        'validation_error_no_elements': 'Si prega di configurare almeno un elemento o selezionare uno stile preimpostato!',
        'validation_error_too_many_references': 'Troppi riferimenti (massimo {} consentiti)',
        'validation_warning_few_references': 'Pochi riferimenti per statistiche significative',
        'cache_initialized': 'Cache inizializzato con successo',
        'cache_cleared': 'Cache cancellato con successo',
        'theme_selector': 'Tema:',
        'light_theme': 'Chiaro',
        'dark_theme': 'Scuro',
        'mobile_view': 'Vista Mobile',
        'desktop_view': 'Vista Desktop',
        'clear_button': '🗑️ Cancella',
        'back_button': '↩️ Indietro'
    },
    'ja': {
        'header': '🎨 引用スタイル構築ツール',
        'general_settings': '⚙️ 一般設定',
        'element_config': '📑 要素設定',
        'style_preview': '👀 スタイルプレビュー',
        'data_input': '📁 データ入力',
        'data_output': '📤 データ出力',
        'numbering_style': '番号付け:',
        'author_format': '著者:',
        'author_separator': '区切り文字:',
        'et_al_limit': 'Et al 以後:',
        'use_and': "'および'",
        'use_ampersand': "'&'",
        'doi_format': 'DOI形式:',
        'doi_hyperlink': 'DOIをハイパーリンクとして',
        'page_format': 'ページ:',
        'final_punctuation': '終了句読点:',
        'element': '要素',
        'italic': '斜体',
        'bold': '太字',
        'parentheses': '括弧',
        'separator': '区切り文字',
        'input_method': '入力:',
        'output_method': '出力:',
        'select_docx': 'DOCXを選択',
        'enter_references': '参考文献を入力（1行に1つ）',
        'references': '参考文献:',
        'results': '結果:',
        'process': '🚀 処理',
        'example': '例:',
        'error_select_element': '少なくとも1つの要素を選択してください！',
        'processing': '⏳ 処理中...',
        'upload_file': 'ファイルをアップロードしてください！',
        'enter_references_error': '参考文献を入力してください！',
        'select_docx_output': 'ダウンロードするDOCX出力を選択してください！',
        'doi_txt': '📄 DOI (TXT)',
        'references_docx': '📋 参考文献 (DOCX)',
        'found_references': '{}件の参考文献が見つかりました。',
        'found_references_text': 'テキスト内で{}件の参考文献が見つかりました。',
        'statistics': '統計: {}件のDOIが見つかりました、{}件は見つかりませんでした。',
        'language': '言語:',
        'gost_style': 'GOSTスタイルを適用',
        'export_style': '📤 スタイルをエクスポート',
        'import_style': '📥 スタイルをインポート',
        'export_file_name': 'ファイル名:',
        'import_file': 'スタイルファイルを選択:',
        'export_success': 'スタイルのエクスポートが成功しました！',
        'import_success': 'スタイルのインポートが成功しました！',
        'import_error': 'スタイルファイルのインポートエラー！',
        'processing_status': '参考文献を処理中...',
        'current_reference': '現在: {}',
        'processed_stats': '処理済み: {}/{} | 見つかった: {} | エラー: {}',
        'time_remaining': '推定残り時間: {}',
        'duplicate_reference': '🔄 重複参考文献 (参照 #{})',
        'batch_processing': 'DOIのバッチ処理...',
        'extracting_metadata': 'メタデータを抽出中...',
        'checking_duplicates': '重複をチェック中...',
        'retrying_failed': '失敗したDOIリクエストを再試行中...',
        'bibliographic_search': '書誌データで検索中...',
        'style_presets': 'スタイルプリセット',
        'gost_button': 'GOST',
        'acs_button': 'ACS (MDPI)',
        'rsc_button': 'RSC',
        'cta_button': 'CTA',
        'style_preset_tooltip': 'ここには、個々の出版社が維持しているいくつかのスタイルがあります。主要な出版社（Elsevier、Springer Nature、Wiley）の場合、スタイルはジャーナルごとに異なります。特定のジャーナルの参考文献を作成（または再フォーマット）するには、引用スタイル構築ツールを使用してください。',
        'journal_style': 'ジャーナルスタイル:',
        'full_journal_name': '完全なジャーナル名',
        'journal_abbr_with_dots': 'J. 略称',
        'journal_abbr_no_dots': 'J 略称',
        'short_guide_title': 'DOIベースの参考文献変換の短いガイド',
        'step_1': '❶ 既製の参考文献スタイル（ACS(MDPI)、RSC、CTA）を選択するか、要素設定の順序、デザイン、句読点を選択して独自のスタイルを作成します',
        'step_1_note': '（！）句読点ボックスを使用すると、要素設定間にさまざまな項目を含めることができます（簡単な句読点、Vol.、Issue…）',
        'step_2': '❷ 次に、スタイルプリセットを使用して、再フォーマットされた各参考文献の特定の要素設定を変更します。',
        'step_3': '❸ スタイルプレビュー機能により、ユーザーは参考文献スタイルの最終的な形を視覚化できます',
        'step_4': '❹ 最終的なスタイルが適切な場合は、データ入力セクションでDocxまたはテキストオプションを選択し、対応する情報をアップロードします。次に、データ輸出セクションで必要なオプションを選択し、「処理」を押して再フォーマットを開始します。',
        'step_5': '❺ 処理が完了した後、希望の形式で再フォーマットされた参考文献をダウンロードします。',
        'step_5_note': '（！）Docxファイルの輸出をお勧めします。これは、フォーマット（太字、斜体、ハイパーリンクなど）を保持し、文書の最後に追加の統計情報を含めるためです。',
        'step_6': '❻ スタイルの最終バージョンを作成した後、次のセッションで再度アップロードできるように保存します。この目的にはスタイル管理セクションを使用してください。',
        'validation_error_no_elements': '少なくとも1つの要素を設定するか、プリセットスタイルを選択してください！',
        'validation_error_too_many_references': '参考文献が多すぎます（最大{}件まで許可）',
        'validation_warning_few_references': '有意な統計のための参考文献が少なすぎます',
        'cache_initialized': 'キャッシュの初期化に成功しました',
        'cache_cleared': 'キャッシュのクリアに成功しました',
        'theme_selector': 'テーマ:',
        'light_theme': 'ライト',
        'dark_theme': 'ダーク',
        'mobile_view': 'モバイル表示',
        'desktop_view': 'デスクトップ表示',
        'clear_button': '🗑️ クリア',
        'back_button': '↩️ 戻る'
    },
    'zh': {
        'header': '🎨 引文样式构建器',
        'general_settings': '⚙️ 通用设置',
        'element_config': '📑 元素配置',
        'style_preview': '👀 样式预览',
        'data_input': '📁 数据输入',
        'data_output': '📤 数据输出',
        'numbering_style': '编号:',
        'author_format': '作者:',
        'author_separator': '分隔符:',
        'et_al_limit': 'Et al 在之后:',
        'use_and': "'和'",
        'use_ampersand': "'&'",
        'doi_format': 'DOI格式:',
        'doi_hyperlink': 'DOI作为超链接',
        'page_format': '页面:',
        'final_punctuation': '结束标点:',
        'element': '元素',
        'italic': '斜体',
        'bold': '粗体',
        'parentheses': '括号',
        'separator': '分隔符',
        'input_method': '输入:',
        'output_method': '输出:',
        'select_docx': '选择DOCX',
        'enter_references': '输入参考文献（每行一个）',
        'references': '参考文献:',
        'results': '结果:',
        'process': '🚀 处理',
        'example': '示例:',
        'error_select_element': '请选择至少一个元素！',
        'processing': '⏳ 处理中...',
        'upload_file': '请上传文件！',
        'enter_references_error': '请输入参考文献！',
        'select_docx_output': '请选择要下载的DOCX输出！',
        'doi_txt': '📄 DOI (TXT)',
        'references_docx': '📋 参考文献 (DOCX)',
        'found_references': '找到 {} 条参考文献。',
        'found_references_text': '在文本中找到 {} 条参考文献。',
        'statistics': '统计: 找到 {} 条DOI，{} 条未找到。',
        'language': '语言:',
        'gost_style': '应用GOST样式',
        'export_style': '📤 导出样式',
        'import_style': '📥 导入样式',
        'export_file_name': '文件名:',
        'import_file': '选择样式文件:',
        'export_success': '样式导出成功！',
        'import_success': '样式导入成功！',
        'import_error': '导入样式文件错误！',
        'processing_status': '处理参考文献中...',
        'current_reference': '当前: {}',
        'processed_stats': '已处理: {}/{} | 找到: {} | 错误: {}',
        'time_remaining': '预计剩余时间: {}',
        'duplicate_reference': '🔄 重复参考文献 (参见 #{})',
        'batch_processing': '批量处理DOI...',
        'extracting_metadata': '提取元数据中...',
        'checking_duplicates': '检查重复项...',
        'retrying_failed': '重试失败的DOI请求...',
        'bibliographic_search': '通过书目数据搜索...',
        'style_presets': '样式预设',
        'gost_button': 'GOST',
        'acs_button': 'ACS (MDPI)',
        'rsc_button': 'RSC',
        'cta_button': 'CTA',
        'style_preset_tooltip': '这里是一些由各个出版商维护的样式。对于主要出版商（Elsevier、Springer Nature、Wiley），样式因期刊而异。要为特定期刊创建（或重新格式化）参考文献，请使用引文样式构建器。',
        'journal_style': '期刊样式:',
        'full_journal_name': '完整期刊名称',
        'journal_abbr_with_dots': '期刊 缩写',
        'journal_abbr_no_dots': '期刊缩写',
        'short_guide_title': '基于DOI的参考文献转换简短指南',
        'step_1': '❶ 选择现成的参考文献样式（ACS(MDPI)、RSC或CTA），或通过选择元素配置的顺序、设计和标点创建自己的样式',
        'step_1_note': '（！）标点框允许在元素配置之间包含各种项目（简单标点、卷、期…）',
        'step_2': '❷ 然后，使用样式预设更改每个重新格式化的参考文献的特定元素配置。',
        'step_3': '❸ 样式预览功能使用户能够可视化其参考文献样式的最终形式',
        'step_4': '❹ 如果最终样式合适，请在数据输入部分选择Docx或文本选项并上传相应信息。然后在数据输出部分选择所需选项并按"处理"开始重新格式化。',
        'step_5': '❺ 处理完成后，以您喜欢的格式下载重新格式化的参考文献。',
        'step_5_note': '（！）建议输出Docx文件，因为它保留格式（例如粗体、斜体和超链接）并在文档末尾包含附加统计信息。',
        'step_6': '❻ 创建样式的最终版本后，保存它以便在下一个会话中再次上传。使用样式管理部分实现此目的。',
        'validation_error_no_elements': '请配置至少一个元素或选择预设样式！',
        'validation_error_too_many_references': '参考文献太多（最多允许 {} 条）',
        'validation_warning_few_references': '参考文献太少，无法生成有意义的统计',
        'cache_initialized': '缓存初始化成功',
        'cache_cleared': '缓存清除成功',
        'theme_selector': '主题:',
        'light_theme': '浅色',
        'dark_theme': '深色',
        'mobile_view': '移动视图',
        'desktop_view': '桌面视图',
        'clear_button': '🗑️ 清除',
        'back_button': '↩️ 返回'
    }
}

# Кэширование DOI
class DOICache:
    """Кэш для хранения метаданных DOI"""
    
    def __init__(self, db_path: str = Config.DB_PATH):
        self.db_path = db_path
        self._init_db()
    
    def _init_db(self):
        """Инициализация базы данных"""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('''
                CREATE TABLE IF NOT EXISTS doi_cache (
                    doi TEXT PRIMARY KEY,
                    metadata TEXT NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    accessed_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_doi ON doi_cache(doi)')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_accessed_at ON doi_cache(accessed_at)')
    
    def get(self, doi: str) -> Optional[Dict]:
        """Получение метаданных из кэша"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                result = conn.execute(
                    'SELECT metadata FROM doi_cache WHERE doi = ? AND datetime(accessed_at) > datetime("now", ?)',
                    (doi, f"-{Config.CACHE_TTL_HOURS} hours")
                ).fetchone()
                
                if result:
                    # Обновляем время доступа
                    conn.execute(
                        'UPDATE doi_cache SET accessed_at = CURRENT_TIMESTAMP WHERE doi = ?',
                        (doi,)
                    )
                    return json.loads(result[0])
        except Exception as e:
            logger.error(f"Cache get error for {doi}: {e}")
        return None
    
    def set(self, doi: str, metadata: Dict):
        """Сохранение метаданных в кэш"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.execute(
                    'INSERT OR REPLACE INTO doi_cache (doi, metadata) VALUES (?, ?)',
                    (doi, json.dumps(metadata))
                )
        except Exception as e:
            logger.error(f"Cache set error for {doi}: {e}")
    
    def clear_old_entries(self):
        """Очистка устаревших записей"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.execute(
                    'DELETE FROM doi_cache WHERE datetime(accessed_at) <= datetime("now", ?)',
                    (f"-{Config.CACHE_TTL_HOURS} hours",)
                )
        except Exception as e:
            logger.error(f"Cache cleanup error: {e}")

# Инициализация кэша
doi_cache = DOICache()

class UserPreferencesManager:
    """Менеджер пользовательских предпочтений"""
    
    def __init__(self, db_path: str = Config.USER_PREFS_DB):
        self.db_path = db_path
        self._init_db()
    
    def _init_db(self):
        """Инициализация базы данных предпочтений"""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('''
                CREATE TABLE IF NOT EXISTS user_preferences (
                    ip_address TEXT PRIMARY KEY,
                    language TEXT DEFAULT 'en',
                    theme TEXT DEFAULT 'light',
                    mobile_view INTEGER DEFAULT 0,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_ip ON user_preferences(ip_address)')
    
    def get_user_ip(self):
        """Получение IP пользователя"""
        try:
            # В Streamlit можно получить IP через экспериментальный API
            if hasattr(st, 'experimental_user'):
                return getattr(st.experimental_user, 'ip', 'unknown')
        except:
            pass
        return 'unknown'
    
    def get_preferences(self, ip: str) -> Dict[str, Any]:
        """Получение предпочтений пользователя"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                result = conn.execute(
                    'SELECT language, theme, mobile_view FROM user_preferences WHERE ip_address = ?',
                    (ip,)
                ).fetchone()
                
                if result:
                    return {
                        'language': result[0],
                        'theme': result[1],
                        'mobile_view': bool(result[2])
                    }
        except Exception as e:
            logger.error(f"Error getting preferences for {ip}: {e}")
        
        return {
            'language': 'en',
            'theme': 'light',
            'mobile_view': False
        }
    
    def save_preferences(self, ip: str, preferences: Dict[str, Any]):
        """Сохранение предпочтений пользователя"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.execute('''
                    INSERT OR REPLACE INTO user_preferences 
                    (ip_address, language, theme, mobile_view, updated_at) 
                    VALUES (?, ?, ?, ?, CURRENT_TIMESTAMP)
                ''', (
                    ip,
                    preferences.get('language', 'en'),
                    preferences.get('theme', 'light'),
                    int(preferences.get('mobile_view', False))
                ))
        except Exception as e:
            logger.error(f"Error saving preferences for {ip}: {e}")
    
    def detect_mobile_device(self, user_agent: str) -> bool:
        """Определение мобильного устройства по User-Agent"""
        try:
            # Простая проверка по ключевым словам в User-Agent
            mobile_keywords = [
                'mobile', 'android', 'iphone', 'ipad', 'tablet', 
                'blackberry', 'webos', 'windows phone'
            ]
            user_agent_lower = user_agent.lower()
            return any(keyword in user_agent_lower for keyword in mobile_keywords)
        except:
            return False

class StyleValidator:
    """Валидатор настроек стиля"""
    
    @staticmethod
    def validate_style_config(style_config: Dict) -> Tuple[bool, List[str]]:
        """Валидация конфигурации стиля"""
        errors = []
        warnings = []
        
        # Проверка наличия элементов или пресетов
        has_elements = bool(style_config.get('elements'))
        has_preset = any([
            style_config.get('gost_style', False),
            style_config.get('acs_style', False), 
            style_config.get('rsc_style', False),
            style_config.get('cta_style', False)
        ])
        
        if not has_elements and not has_preset:
            errors.append(get_text('validation_error_no_elements'))
        
        # Проверка корректности элементов
        if has_elements:
            elements = style_config['elements']
            for i, (element, config) in enumerate(elements):
                if not element:
                    errors.append(f"Element {i+1} is empty")
                if not config.get('separator', '').strip() and i < len(elements) - 1:
                    warnings.append(f"Element {i+1} has empty separator")
        
        return len(errors) == 0, errors + warnings
    
    @staticmethod
    def validate_references_count(references: List[str]) -> Tuple[bool, List[str]]:
        """Валидация количества ссылок"""
        errors = []
        warnings = []
        
        if len(references) > Config.MAX_REFERENCES:
            errors.append(get_text('validation_error_too_many_references').format(Config.MAX_REFERENCES))
        
        if len(references) < Config.MIN_REFERENCES_FOR_STATS:
            warnings.append(get_text('validation_warning_few_references'))
        
        return len(errors) == 0, errors + warnings

class ProgressManager:
    """Менеджер прогресса обработки"""
    
    def __init__(self):
        self.start_time = None
        self.progress_data = {
            'total': 0,
            'processed': 0,
            'found': 0,
            'errors': 0,
            'phase': 'initializing'
        }
    
    def start_processing(self, total: int):
        """Начало обработки"""
        self.start_time = time.time()
        self.progress_data = {
            'total': total,
            'processed': 0,
            'found': 0,
            'errors': 0,
            'phase': 'processing'
        }
    
    def update_progress(self, processed: int, found: int, errors: int, phase: str = None):
        """Обновление прогресса"""
        self.progress_data.update({
            'processed': processed,
            'found': found,
            'errors': errors
        })
        if phase:
            self.progress_data['phase'] = phase
    
    def get_progress_info(self) -> Dict[str, Any]:
        """Получение информации о прогрессе"""
        if not self.start_time:
            return self.progress_data
        
        elapsed = time.time() - self.start_time
        processed = self.progress_data['processed']
        total = self.progress_data['total']
        
        # Расчет оставшегося времени
        time_remaining = None
        if processed > 0 and total > 0:
            estimated_total = (elapsed / processed) * total
            time_remaining = estimated_total - elapsed
            if time_remaining < 0:
                time_remaining = 0
        
        # Расчет прогресса для цветового градиента
        progress_ratio = processed / total if total > 0 else 0
        
        return {
            **self.progress_data,
            'elapsed_time': elapsed,
            'time_remaining': time_remaining,
            'progress_ratio': progress_ratio
        }
    
    def get_progress_color(self, progress_ratio: float) -> str:
        """Получение цвета прогресс-бара на основе прогресса"""
        if progress_ratio < 0.33:
            return Config.PROGRESS_COLORS['start']
        elif progress_ratio < 0.66:
            return Config.PROGRESS_COLORS['middle']
        else:
            return Config.PROGRESS_COLORS['end']

# Инициализация глобальных состояний
def init_session_state():
    """Инициализация состояния сессии"""
    defaults = {
        'current_language': 'en',
        'current_theme': 'light',
        'mobile_view': False,
        'imported_style': None,
        'style_applied': False,
        'apply_imported_style': False,
        'output_text_value': "",
        'show_results': False,
        'download_data': {},
        'use_and_checkbox': False,
        'use_ampersand_checkbox': False,
        'journal_style': '{Full Journal Name}',
        'num': "No numbering",
        'auth': "AA Smith",
        'sep': ", ",
        'etal': 0,
        'doi': "10.10/xxx",
        'doilink': True,
        'page': "122–128",
        'punct': "",
        'gost_style': False,
        'acs_style': False,
        'rsc_style': False,
        'cta_style': False,
        'last_style_update': 0,
        'cache_initialized': False,
        'user_prefs_loaded': False,
        'file_processing_complete': False,
        'style_import_processed': False,  # Флаг для отслеживания обработки импорта
        'last_imported_file_hash': None,  # Хеш последнего импортированного файла
        'style_management_initialized': False,  # Флаг инициализации управления стилями
        'previous_states': [],  # Стек предыдущих состояний для кнопки Back
        'max_undo_steps': 10,  # Максимальное количество шагов отмены
    }
    
    for key, default in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = default
    
    # Инициализация элементов конфигурации
    for i in range(8):
        for prop in ['el', 'it', 'bd', 'pr', 'sp']:
            key = f"{prop}{i}"
            if key not in st.session_state:
                if prop == 'sp':
                    st.session_state[key] = ". "
                elif prop == 'el':
                    st.session_state[key] = ""
                else:
                    st.session_state[key] = False

def get_text(key: str) -> str:
    """Получение перевода по ключу"""
    return TRANSLATIONS[st.session_state.current_language].get(key, key)

# Базовые классы форматирования
class JournalAbbreviation:
    def __init__(self):
        self.ltwa_data = {}
        self.load_ltwa_data()
        self.uppercase_abbreviations = {'acs', 'ecs', 'rsc', 'ieee', 'iet', 'acm', 'aims', 'bmc', 'bmj', 'npj'}
        self.special_endings = {'A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 
                               'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z',
                               'I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X'}
    
    def load_ltwa_data(self):
        """Загружает данные сокращений из файла ltwa.csv"""
        try:
            csv_path = Config.LTWA_CSV_PATH
            if os.path.exists(csv_path):
                with open(csv_path, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f, delimiter='\t')
                    next(reader)
                    for row in reader:
                        if len(row) >= 2:
                            word = row[0].strip()
                            abbreviation = row[1].strip() if row[1].strip() else None
                            self.ltwa_data[word] = abbreviation
            else:
                logger.warning(f"Файл {csv_path} не найден, используется стандартное сокращение")
        except Exception as e:
            logger.error(f"Ошибка загрузки ltwa.csv: {e}")
    
    def abbreviate_word(self, word: str) -> str:
        """Сокращает одно слово на основе данных LTWA"""
        word_lower = word.lower()
        
        if word_lower in self.ltwa_data:
            abbr = self.ltwa_data[word_lower]
            return abbr if abbr else word
        
        for ltwa_word, abbr in self.ltwa_data.items():
            if ltwa_word.endswith('-') and word_lower.startswith(ltwa_word[:-1]):
                return abbr if abbr else word
        
        return word
    
    def extract_special_endings(self, journal_name: str) -> Tuple[str, str]:
        """Извлекает специальные окончания (A, B, C и т.д.) из названия журнала"""
        # Паттерны для поиска специальных окончаний
        patterns = [
            r'\s+([A-Z])\s*$',  # Одиночные буквы в конце
            r'\s+([IVX]+)\s*$',  # Римские цифры
            r'\s+Part\s+([A-Z0-9]+)\s*$',  # Part A, Part 1 и т.д.
            r'\s+([A-Z]):\s+[A-Z]',  # Буква с двоеточием: A: General, B: Environmental
        ]
        
        for pattern in patterns:
            match = re.search(pattern, journal_name)
            if match:
                ending = match.group(1)
                # Проверяем, является ли окончание специальным
                if ending in self.special_endings or re.match(r'^[A-Z]$', ending):
                    base_name = journal_name[:match.start()].strip()
                    return base_name, ending
        
        return journal_name, ""
    
    def abbreviate_journal_name(self, journal_name: str, style: str = "{J. Abbr.}") -> str:
        """Сокращает название журнала в соответствии с выбранным стилем"""
        if not journal_name:
            return ""
        
        # Извлекаем базовое название и специальное окончание
        base_name, special_ending = self.extract_special_endings(journal_name)
        
        words_to_remove = {'a', 'an', 'the', 'of', 'in', 'and', '&', 'for', 'on', 'with', 'by'}
        words = [word for word in base_name.split() if word.lower() not in words_to_remove]
        words = [word.replace(':', '') for word in words]
        
        if len(words) <= 1:
            result = journal_name
        else:
            abbreviated_words = []
            for i, word in enumerate(words):
                original_first_char = word[0]
                abbreviated = self.abbreviate_word(word.lower())
                
                if abbreviated and original_first_char.isupper():
                    abbreviated = abbreviated[0].upper() + abbreviated[1:]
                
                if i == 0 and abbreviated.lower() in self.uppercase_abbreviations:
                    abbreviated = abbreviated.upper()
                
                abbreviated_words.append(abbreviated)
            
            if style == "{J. Abbr.}":
                result = " ".join(abbreviated_words)
            elif style == "{J Abbr}":
                result = " ".join(abbr.replace('.', '') for abbr in abbreviated_words)
            else:
                result = base_name
        
        # Добавляем специальное окончание обратно
        if special_ending:
            if ':' in journal_name and special_ending + ':' in journal_name:
                # Для случаев типа "Applied Catalysis A: General"
                result += f" {special_ending}:"
                # Добавляем остаток после двоеточия
                after_colon = journal_name.split(special_ending + ':', 1)[1].strip()
                if after_colon:
                    result += f" {after_colon}"
            else:
                result += f" {special_ending}"
        
        result = re.sub(r'\.\.+', '.', result)
        return result

# Инициализация системы сокращений
journal_abbrev = JournalAbbreviation()

class BaseCitationFormatter:
    """Базовый класс для форматирования цитирования"""
    
    def __init__(self, style_config: Dict[str, Any]):
        self.style_config = style_config
    
    def format_authors(self, authors: List[Dict[str, str]]) -> str:
        """Форматирует список авторов"""
        if not authors:
            return ""
        
        author_format = self.style_config['author_format']
        separator = self.style_config['author_separator']
        et_al_limit = self.style_config['et_al_limit']
        use_and_bool = self.style_config['use_and_bool']
        use_ampersand_bool = self.style_config['use_ampersand_bool']
        
        author_str = ""
        
        if use_and_bool or use_ampersand_bool:
            limit = len(authors)
        else:
            limit = et_al_limit if et_al_limit and et_al_limit > 0 else len(authors)
        
        for i, author in enumerate(authors[:limit]):
            given = author['given']
            family = author['family']
            
            initials = given.split()[:2]
            first_initial = initials[0][0] if initials else ''
            second_initial = initials[1][0].upper() if len(initials) > 1 else ''
            
            if author_format == "AA Smith":
                formatted_author = f"{first_initial}{second_initial} {family}"
            elif author_format == "A.A. Smith":
                if second_initial:
                    formatted_author = f"{first_initial}.{second_initial}. {family}"
                else:
                    formatted_author = f"{first_initial}. {family}"
            elif author_format == "Smith AA":
                formatted_author = f"{family} {first_initial}{second_initial}"
            elif author_format == "Smith A.A":
                if second_initial:
                    formatted_author = f"{family} {first_initial}.{second_initial}."
                else:
                    formatted_author = f"{family} {first_initial}."
            elif author_format == "Smith, A.A.":
                if second_initial:
                    formatted_author = f"{family}, {first_initial}.{second_initial}."
                else:
                    formatted_author = f"{family}, {first_initial}."
            else:
                formatted_author = f"{first_initial}. {family}"
            
            author_str += formatted_author
            
            if i < len(authors[:limit]) - 1:
                if i == len(authors[:limit]) - 2 and (use_and_bool or use_ampersand_bool):
                    if use_and_bool:
                        author_str += " and "
                    else:
                        author_str += " & "
                else:
                    author_str += separator
        
        if et_al_limit and len(authors) > et_al_limit and not (use_and_bool or use_ampersand_bool):
            author_str += " et al"
        
        return author_str.strip()
    
    def format_pages(self, pages: str, article_number: str, style_type: str = "default") -> str:
        """Форматирует страницы в зависимости от стиля"""
        page_format = self.style_config['page_format']
        
        if pages:
            if style_type == "rsc":
                if '-' in pages:
                    first_page = pages.split('-')[0].strip()
                    return first_page
                else:
                    return pages.strip()
            elif style_type == "cta":
                if '-' in pages:
                    start, end = pages.split('-')
                    start = start.strip()
                    end = end.strip()
                    
                    if len(start) == len(end) and start[:-1] == end[:-1]:
                        return f"{start}–{end[-1]}"
                    elif len(start) > 1 and len(end) > 1 and start[:-2] == end[:-2]:
                        return f"{start}–{end[-2:]}"
                    else:
                        return f"{start}–{end}"
                else:
                    return pages.strip()
            else:
                if '-' not in pages:
                    return pages
                
                start, end = pages.split('-')
                start = start.strip()
                end = end.strip()
                
                if page_format == "122 - 128":
                    return f"{start} - {end}"
                elif page_format == "122-128":
                    return f"{start}-{end}"
                elif page_format == "122 – 128":
                    return f"{start} – {end}"
                elif page_format == "122–128":
                    return f"{start}–{end}"
                elif page_format == "122–8":
                    i = 0
                    while i < len(start) and i < len(end) and start[i] == end[i]:
                        i += 1
                    return f"{start}–{end[i:]}"
        
        return article_number
    
    def format_doi(self, doi: str) -> Tuple[str, str]:
        """Форматирует DOI и возвращает текст и URL"""
        doi_format = self.style_config['doi_format']
        
        if doi_format == "10.10/xxx":
            value = doi
        elif doi_format == "doi:10.10/xxx":
            value = f"doi:{doi}"
        elif doi_format == "DOI:10.10/xxx":
            value = f"DOI:{doi}"
        elif doi_format == "https://dx.doi.org/10.10/xxx":
            value = f"https://dx.doi.org/{doi}"
        else:
            value = doi
        
        return value, f"https://doi.org/{doi}"
    
    def format_journal_name(self, journal_name: str) -> str:
        """Форматирует название журнала с учетом выбранного стиля"""
        journal_style = self.style_config.get('journal_style', '{Full Journal Name}')
        return journal_abbrev.abbreviate_journal_name(journal_name, journal_style)

class CustomCitationFormatter(BaseCitationFormatter):
    """Форматировщик для пользовательских стилей с улучшенной обработкой Issue"""
    
    def format_reference(self, metadata: Dict[str, Any], for_preview: bool = False) -> Tuple[Any, bool]:
        if not metadata:
            error_message = "Ошибка: Не удалось отформатировать ссылку." if st.session_state.current_language == 'ru' else "Error: Could not format the reference."
            return (error_message, True)
        
        elements = []
        previous_element_was_empty = False
        
        for i, (element, config) in enumerate(self.style_config['elements']):
            value = ""
            doi_value = None
            element_empty = False
            
            if element == "Authors":
                value = self.format_authors(metadata['authors'])
                element_empty = not value
            elif element == "Title":
                value = metadata['title']
                element_empty = not value
            elif element == "Journal":
                value = self.format_journal_name(metadata['journal'])
                element_empty = not value
            elif element == "Year":
                value = str(metadata['year']) if metadata['year'] else ""
                element_empty = not value
            elif element == "Volume":
                value = metadata['volume']
                element_empty = not value
            elif element == "Issue":
                value = metadata['issue']
                element_empty = not value
            elif element == "Pages":
                value = self.format_pages(metadata['pages'], metadata['article_number'])
                element_empty = not value
            elif element == "DOI":
                doi = metadata['doi']
                doi_value = doi
                value, _ = self.format_doi(doi)
                element_empty = not value
            
            # Обработка пустых элементов и их разделителей
            if value:
                if config['parentheses'] and value:
                    value = f"({value})"
                
                # Определяем разделитель с учетом пустых элементов
                separator = ""
                if i < len(self.style_config['elements']) - 1:
                    if not element_empty:
                        # Если текущий элемент не пустой, используем его разделитель
                        separator = config['separator']
                    elif previous_element_was_empty:
                        # Если предыдущий элемент был пустой, пропускаем разделитель
                        separator = ""
                    else:
                        # Если текущий элемент пустой, но предыдущий был не пустой, используем разделитель
                        separator = config['separator']
                
                if for_preview:
                    formatted_value = value
                    if config['italic']:
                        formatted_value = f"<i>{formatted_value}</i>"
                    if config['bold']:
                        formatted_value = f"<b>{formatted_value}</b>"
                    
                    elements.append((formatted_value, False, False, separator, False, None, element_empty))
                else:
                    elements.append((value, config['italic'], config['bold'], separator,
                                   (element == "DOI" and self.style_config['doi_hyperlink']), doi_value, element_empty))
                
                previous_element_was_empty = False
            else:
                # Элемент пустой - запоминаем это для следующей итерации
                previous_element_was_empty = True
        
        # Пост-обработка для удаления лишних разделителей
        cleaned_elements = []
        for i, element_data in enumerate(elements):
            value, italic, bold, separator, is_doi_hyperlink, doi_value, element_empty = element_data
            
            # Если элемент не пустой, добавляем его
            if not element_empty:
                # Для последнего элемента убираем разделитель
                if i == len(elements) - 1:
                    separator = ""
                
                cleaned_elements.append((value, italic, bold, separator, is_doi_hyperlink, doi_value))
        
        if for_preview:
            ref_str = ""
            for i, (value, _, _, separator, _, _) in enumerate(cleaned_elements):
                ref_str += value
                if separator and i < len(cleaned_elements) - 1:
                    ref_str += separator
                elif i == len(cleaned_elements) - 1 and self.style_config['final_punctuation']:
                    ref_str = ref_str.rstrip(',.') + "."
            
            ref_str = re.sub(r'\.\.+', '.', ref_str)
            return ref_str, False
        else:
            return cleaned_elements, False

class GOSTCitationFormatter(BaseCitationFormatter):
    """Форматировщик для стиля ГОСТ (обновленная версия)"""
    
    def format_reference(self, metadata: Dict[str, Any], for_preview: bool = False) -> Tuple[Any, bool]:
        if not metadata:
            error_message = "Ошибка: Не удалось отформатировать ссылку." if st.session_state.current_language == 'ru' else "Error: Could not format the reference."
            return (error_message, True)
        
        # Форматирование авторов в новом формате: Smith J.A., Doe A.B.
        authors_str = ""
        for i, author in enumerate(metadata['authors']):
            given = author['given']
            family = author['family']
            initials = given.split()[:2]
            first_initial = initials[0][0] if initials else ''
            second_initial = initials[1][0].upper() if len(initials) > 1 else ''
            
            if second_initial:
                author_str = f"{family} {first_initial}.{second_initial}."
            else:
                author_str = f"{family} {first_initial}."
            
            authors_str += author_str
            
            if i < len(metadata['authors']) - 1:
                authors_str += ", "
        
        pages = metadata['pages']
        article_number = metadata['article_number']
        
        # Используем полное название журнала
        journal_name = metadata['journal']
        
        doi_url = f"https://doi.org/{metadata['doi']}"
        
        # Форматирование основной ссылки
        if metadata['issue']:
            gost_ref = f"{authors_str} {metadata['title']} // {journal_name}. – {metadata['year']}. – Vol. {metadata['volume']}, № {metadata['issue']}"
        else:
            gost_ref = f"{authors_str} {metadata['title']} // {journal_name}. – {metadata['year']}. – Vol. {metadata['volume']}"
        
        # НОВАЯ ЛОГИКА: Приоритет article-number над pages
        if article_number and article_number.strip():
            # Используем номер статьи (высший приоритет)
            gost_ref += f". – Art. {article_number.strip()}"
        elif pages and pages.strip():
            # Используем страницы (если нет article-number)
            # Форматирование страниц в формате "122-128" (с обычным дефисом)
            if '-' in pages:
                start_page, end_page = pages.split('-')
                pages_formatted = f"{start_page.strip()}-{end_page.strip()}"
            else:
                pages_formatted = pages.strip()
            gost_ref += f". – Р. {pages_formatted}"
        else:
            # Нет ни article-number, ни pages
            if st.session_state.current_language == 'ru':
                gost_ref += ". – [Без пагинации]"
            else:
                gost_ref += ". – [No pagination]"
        
        # Добавляем DOI
        gost_ref += f". – {doi_url}"
        
        if for_preview:
            return gost_ref, False
        else:
            elements = []
            text_before_doi = gost_ref.replace(doi_url, "")
            elements.append((text_before_doi, False, False, "", False, None))
            elements.append((doi_url, False, False, "", True, metadata['doi']))
            return elements, False

class ACSCitationFormatter(BaseCitationFormatter):
    """Форматировщик для стиля ACS (MDPI)"""
    
    def format_reference(self, metadata: Dict[str, Any], for_preview: bool = False) -> Tuple[Any, bool]:
        if not metadata:
            error_message = "Ошибка: Не удалось отформатировать ссылку." if st.session_state.current_language == 'ru' else "Error: Could not format the reference."
            return (error_message, True)
        
        authors_str = ""
        for i, author in enumerate(metadata['authors']):
            given = author['given']
            family = author['family']
            
            initials = given.split()[:2]
            first_initial = initials[0][0] if initials else ''
            second_initial = initials[1][0].upper() if len(initials) > 1 else ''
            
            if second_initial:
                author_str = f"{family}, {first_initial}.{second_initial}."
            else:
                author_str = f"{family}, {first_initial}."
            
            authors_str += author_str
            
            if i < len(metadata['authors']) - 1:
                authors_str += "; "
        
        pages = metadata['pages']
        article_number = metadata['article_number']
        
        # ИЗМЕНЕНИЕ 1: Используем полный формат страниц вместо сокращенного
        if pages:
            if '-' in pages:
                start_page, end_page = pages.split('-')
                start_page = start_page.strip()
                end_page = end_page.strip()
                # Убираем сокращение и используем полный формат
                pages_formatted = f"{start_page}–{end_page}"
            else:
                pages_formatted = pages
        elif article_number:
            pages_formatted = article_number
        else:
            pages_formatted = ""
        
        journal_name = self.format_journal_name(metadata['journal'])
        
        # Форматируем DOI как гиперссылку
        doi_url = f"https://dx.doi.org/{metadata['doi']}"
        
        # ИЗМЕНЕНИЕ 2: Добавляем DOI после страниц через ". "
        acs_ref = f"{authors_str} {metadata['title']}. {journal_name} {metadata['year']}, {metadata['volume']}, {pages_formatted}. {doi_url}"
        acs_ref = re.sub(r'\.\.+', '.', acs_ref)
        
        if for_preview:
            return acs_ref, False
        else:
            elements = []
            elements.append((authors_str, False, False, " ", False, None))
            elements.append((metadata['title'], False, False, ". ", False, None))
            elements.append((journal_name, True, False, " ", False, None))
            elements.append((str(metadata['year']), False, True, ", ", False, None))
            elements.append((metadata['volume'], True, False, ", ", False, None))
            elements.append((pages_formatted, False, False, ". ", False, None))
            # ИЗМЕНЕНИЕ 3: Добавляем DOI как отдельный элемент с гиперссылкой
            elements.append((doi_url, False, False, "", True, metadata['doi']))
            return elements, False

class RSCCitationFormatter(BaseCitationFormatter):
    """Форматировщик для стиля RSC"""
    
    def format_reference(self, metadata: Dict[str, Any], for_preview: bool = False) -> Tuple[Any, bool]:
        if not metadata:
            error_message = "Ошибка: Не удалось отформатировать ссылку." if st.session_state.current_language == 'ru' else "Error: Could not format the reference."
            return (error_message, True)
        
        authors_str = ""
        for i, author in enumerate(metadata['authors']):
            given = author['given']
            family = author['family']
            
            initials = given.split()[:2]
            first_initial = initials[0][0] if initials else ''
            second_initial = initials[1][0].upper() if len(initials) > 1 else ''
            
            if second_initial:
                author_str = f"{first_initial}.{second_initial}. {family}"
            else:
                author_str = f"{first_initial}. {family}"
            
            authors_str += author_str
            
            if i < len(metadata['authors']) - 1:
                if i == len(metadata['authors']) - 2:
                    authors_str += " and "
                else:
                    authors_str += ", "
        
        pages = metadata['pages']
        article_number = metadata['article_number']
        
        if pages:
            if '-' in pages:
                first_page = pages.split('-')[0].strip()
                pages_formatted = first_page
            else:
                pages_formatted = pages.strip()
        elif article_number:
            pages_formatted = article_number
        else:
            pages_formatted = ""
        
        journal_name = self.format_journal_name(metadata['journal'])
        rsc_ref = f"{authors_str}, {journal_name}, {metadata['year']}, {metadata['volume']}, {pages_formatted}."
        rsc_ref = re.sub(r'\.\.+', '.', rsc_ref)
        
        if for_preview:
            return rsc_ref, False
        else:
            elements = []
            elements.append((authors_str, False, False, ", ", False, None))
            elements.append((journal_name, True, False, ", ", False, None))
            elements.append((str(metadata['year']), False, False, ", ", False, None))
            elements.append((metadata['volume'], False, True, ", ", False, None))
            elements.append((pages_formatted, False, False, ".", False, None))
            return elements, False

class CTACitationFormatter(BaseCitationFormatter):
    """Форматировщик для стиля CTA"""
    
    def format_reference(self, metadata: Dict[str, Any], for_preview: bool = False) -> Tuple[Any, bool]:
        if not metadata:
            error_message = "Ошибка: Не удалось отформатировать ссылку." if st.session_state.current_language == 'ru' else "Error: Could not format the reference."
            return (error_message, True)
        
        authors_str = ""
        for i, author in enumerate(metadata['authors']):
            given = author['given']
            family = author['family']
            
            initials = given.split()[:2]
            first_initial = initials[0][0] if initials else ''
            second_initial = initials[1][0].upper() if len(initials) > 1 else ''
            
            if second_initial:
                author_str = f"{family} {first_initial}{second_initial}"
            else:
                author_str = f"{family} {first_initial}"
            
            authors_str += author_str
            
            if i < len(metadata['authors']) - 1:
                authors_str += ", "
        
        pages = metadata['pages']
        article_number = metadata['article_number']
        pages_formatted = self.format_pages(pages, article_number, "cta")
        journal_name = self.format_journal_name(metadata['journal'])
        issue_part = f"({metadata['issue']})" if metadata['issue'] else ""
        
        cta_ref = f"{authors_str}. {metadata['title']}. {journal_name}. {metadata['year']};{metadata['volume']}{issue_part}:{pages_formatted}. doi:{metadata['doi']}"
        
        if for_preview:
            return cta_ref, False
        else:
            elements = []
            elements.append((authors_str, False, False, ". ", False, None))
            elements.append((metadata['title'], False, False, ". ", False, None))
            elements.append((journal_name, True, False, ". ", False, None))
            elements.append((str(metadata['year']), False, False, ";", False, None))
            elements.append((metadata['volume'], False, False, "", False, None))
            if metadata['issue']:
                elements.append((f"({metadata['issue']})", False, False, ":", False, None))
            else:
                elements.append(("", False, False, ":", False, None))
            elements.append((pages_formatted, False, False, ". ", False, None))
            doi_text = f"doi:{metadata['doi']}"
            elements.append((doi_text, False, False, "", True, metadata['doi']))
            return elements, False

class CitationFormatterFactory:
    """Фабрика для создания форматировщиков цитирования"""
    
    @staticmethod
    def create_formatter(style_config: Dict[str, Any]) -> BaseCitationFormatter:
        if style_config.get('gost_style', False):
            return GOSTCitationFormatter(style_config)
        elif style_config.get('acs_style', False):
            return ACSCitationFormatter(style_config)
        elif style_config.get('rsc_style', False):
            return RSCCitationFormatter(style_config)
        elif style_config.get('cta_style', False):
            return CTACitationFormatter(style_config)
        else:
            return CustomCitationFormatter(style_config)

class DocumentGenerator:
    """Класс для генерации DOCX документов"""
    
    @staticmethod
    def add_hyperlink(paragraph, text, url):
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        rPr.append(underline)
        
        new_run.append(rPr)
        new_text = OxmlElement('w:t')
        new_text.text = text
        new_run.append(new_text)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        
        return hyperlink
    
    @staticmethod
    def apply_yellow_background(run):
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'FFFF00')
        run._element.get_or_add_rPr().append(shd)
    
    @staticmethod
    def apply_blue_background(run):
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'E6F3FF')
        run._element.get_or_add_rPr().append(shd)
    
    @staticmethod
    def apply_red_color(run):
        color = OxmlElement('w:color')
        color.set(qn('w:val'), 'FF0000')
        run._element.get_or_add_rPr().append(color)
    
    @staticmethod
    def generate_document(formatted_refs: List[Tuple[Any, bool, Any]], 
                         statistics: Dict[str, Any],
                         style_config: Dict[str, Any],
                         duplicates_info: Dict[int, int] = None) -> io.BytesIO:
        output_doc = Document()
        output_doc.add_paragraph('Citation Style Construction / © IHTE, https://ihte.ru/ © CTA, https://chimicatechnoacta.ru / developed by daM©')
        output_doc.add_paragraph('See short stats after the References section')
        output_doc.add_heading('References', level=1)
        
        DocumentGenerator._add_formatted_references(output_doc, formatted_refs, style_config, duplicates_info)
        DocumentGenerator._add_statistics_section(output_doc, statistics)
        
        output_doc_buffer = io.BytesIO()
        output_doc.save(output_doc_buffer)
        output_doc_buffer.seek(0)
        return output_doc_buffer
    
    @staticmethod
    def _add_formatted_references(doc: Document, 
                                formatted_refs: List[Tuple[Any, bool, Any]], 
                                style_config: Dict[str, Any],
                                duplicates_info: Dict[int, int] = None):
        for i, (elements, is_error, metadata) in enumerate(formatted_refs):
            numbering = style_config['numbering_style']
            
            if numbering == "No numbering":
                prefix = ""
            elif numbering == "1":
                prefix = f"{i + 1} "
            elif numbering == "1.":
                prefix = f"{i + 1}. "
            elif numbering == "1)":
                prefix = f"{i + 1}) "
            elif numbering == "(1)":
                prefix = f"({i + 1}) "
            elif numbering == "[1]":
                prefix = f"[{i + 1}] "
            else:
                prefix = f"{i + 1}. "
            
            para = doc.add_paragraph(prefix)
            
            if is_error:
                run = para.add_run(str(elements))
                DocumentGenerator.apply_yellow_background(run)
            elif duplicates_info and i in duplicates_info:
                original_index = duplicates_info[i] + 1
                duplicate_note = get_text('duplicate_reference').format(original_index)
                
                if isinstance(elements, str):
                    run = para.add_run(elements)
                    DocumentGenerator.apply_blue_background(run)
                    para.add_run(f" - {duplicate_note}").italic = True
                else:
                    for j, (value, italic, bold, separator, is_doi_hyperlink, doi_value) in enumerate(elements):
                        if is_doi_hyperlink and doi_value:
                            DocumentGenerator.add_hyperlink(para, value, f"https://doi.org/{doi_value}")
                        else:
                            run = para.add_run(value)
                            if italic:
                                run.font.italic = True
                            if bold:
                                run.font.bold = True
                            DocumentGenerator.apply_blue_background(run)
                        
                        if separator and j < len(elements) - 1:
                            para.add_run(separator)
                    
                    para.add_run(f" - {duplicate_note}").italic = True
            else:
                if metadata is None:
                    run = para.add_run(str(elements))
                    run.font.italic = True
                else:
                    for j, (value, italic, bold, separator, is_doi_hyperlink, doi_value) in enumerate(elements):
                        if is_doi_hyperlink and doi_value:
                            DocumentGenerator.add_hyperlink(para, value, f"https://doi.org/{doi_value}")
                        else:
                            run = para.add_run(value)
                            if italic:
                                run.font.italic = True
                            if bold:
                                run.font.bold = True
                        
                        if separator and j < len(elements) - 1:
                            para.add_run(separator)
                    
                    if style_config['final_punctuation'] and not is_error:
                        para.add_run(".")
    
    @staticmethod
    def _add_statistics_section(doc: Document, statistics: Dict[str, Any]):
        doc.add_heading('Stats', level=1)
        
        doc.add_heading('Journal Frequency', level=2)
        journal_table = doc.add_table(rows=1, cols=3)
        journal_table.style = 'Table Grid'
        
        hdr_cells = journal_table.rows[0].cells
        hdr_cells[0].text = 'Journal Name'
        hdr_cells[1].text = 'Count'
        hdr_cells[2].text = 'Percentage (%)'
        
        for journal_stat in statistics['journal_stats']:
            row_cells = journal_table.add_row().cells
            row_cells[0].text = journal_stat['journal']
            row_cells[1].text = str(journal_stat['count'])
            row_cells[2].text = str(journal_stat['percentage'])
        
        doc.add_paragraph()
        
        doc.add_heading('Year Distribution', level=2)
        
        if statistics['needs_more_recent_references']:
            warning_para = doc.add_paragraph()
            warning_run = warning_para.add_run("To improve the relevance and significance of the research, consider including more recent references published within the last 3-4 years")
            DocumentGenerator.apply_red_color(warning_run)
            doc.add_paragraph()
        
        year_table = doc.add_table(rows=1, cols=3)
        year_table.style = 'Table Grid'
        
        hdr_cells = year_table.rows[0].cells
        hdr_cells[0].text = 'Year'
        hdr_cells[1].text = 'Count'
        hdr_cells[2].text = 'Percentage (%)'
        
        for year_stat in statistics['year_stats']:
            row_cells = year_table.add_row().cells
            row_cells[0].text = str(year_stat['year'])
            row_cells[1].text = str(year_stat['count'])
            row_cells[2].text = str(year_stat['percentage'])
        
        doc.add_paragraph()
        
        doc.add_heading('Author Distribution', level=2)
        
        if statistics['has_frequent_author']:
            warning_para = doc.add_paragraph()
            warning_run = warning_para.add_run("The author(s) are referenced frequently. Either reduce the number of references to the author(s), or expand the reference list to include more sources")
            DocumentGenerator.apply_red_color(warning_run)
            doc.add_paragraph()
        
        author_table = doc.add_table(rows=1, cols=3)
        author_table.style = 'Table Grid'
        
        hdr_cells = author_table.rows[0].cells
        hdr_cells[0].text = 'Author'
        hdr_cells[1].text = 'Count'
        hdr_cells[2].text = 'Percentage (%)'
        
        for author_stat in statistics['author_stats']:
            row_cells = author_table.add_row().cells
            row_cells[0].text = author_stat['author']
            row_cells[1].text = str(author_stat['count'])
            row_cells[2].text = str(author_stat['percentage'])

# Улучшенные функции обработки DOI
class DOIProcessor:
    """Процессор для работы с DOI"""
    
    def __init__(self):
        self.cache = doi_cache
        self.works = works
    
    def find_doi_enhanced(self, reference: str) -> Optional[str]:
        """Улучшенный поиск DOI с использованием нескольких стратегий"""
        if self._is_section_header(reference):
            return None
        
        # Стратегия 1: Поиск явного DOI
        explicit_doi = self._find_explicit_doi(reference)
        if explicit_doi:
            logger.info(f"Found explicit DOI: {explicit_doi}")
            return explicit_doi
        
        # Стратегия 2: Поиск по библиографическим данным в Crossref
        bibliographic_doi = self._find_bibliographic_doi(reference)
        if bibliographic_doi:
            logger.info(f"Found bibliographic DOI: {bibliographic_doi}")
            return bibliographic_doi
        
        # Стратегия 3: Поиск через OpenAlex (если подключен)
        openalex_doi = self._find_openalex_doi(reference)
        if openalex_doi:
            logger.info(f"Found OpenAlex DOI: {openalex_doi}")
            return openalex_doi
        
        logger.warning(f"No DOI found for reference: {reference[:100]}...")
        return None
    
    def _is_section_header(self, text: str) -> bool:
        """Определяет, является ли текст заголовком раздела"""
        text_upper = text.upper().strip()
        section_patterns = [
            r'^NOTES?\s+AND\s+REFERENCES?$',
            r'^REFERENCES?$',
            r'^BIBLIOGRAPHY$',
            r'^LITERATURE$',
            r'^WORKS?\s+CITED$',
            r'^SOURCES?$',
            r'^CHAPTER\s+\d+$',
            r'^SECTION\s+\d+$',
            r'^PART\s+\d+$'
        ]
        
        for pattern in section_patterns:
            if re.search(pattern, text_upper):
                return True
        return False
    
    def _find_explicit_doi(self, reference: str) -> Optional[str]:
        """Поиск явного DOI в тексте"""
        doi_patterns = [
            r'https?://doi\.org/(10\.\d{4,9}/[-._;()/:A-Za-z0-9]+)',
            r'doi:\s*(10\.\d{4,9}/[-._;()/:A-Za-z0-9]+)',
            r'DOI:\s*(10\.\d{4,9}/[-._;()/:A-Za-z0-9]+)',
            r'\b(10\.\d{4,9}/[-._;()/:A-Za-z0-9]+)\b'
        ]
        
        for pattern in doi_patterns:
            match = re.search(pattern, reference, re.IGNORECASE)
            if match:
                doi = match.group(1).rstrip('.,;:')
                return doi
        
        clean_ref = reference.strip()
        if re.match(r'^(doi:|DOI:)?\s*10\.\d{4,9}/[-._;()/:A-Za-z0-9]+\s*$', clean_ref, re.IGNORECASE):
            doi_match = re.search(r'(10\.\d{4,9}/[-._;()/:A-Za-z0-9]+)', clean_ref)
            if doi_match:
                return doi_match.group(1).rstrip('.,;:')
        
        return None
    
    def _find_bibliographic_doi(self, reference: str) -> Optional[str]:
        """Поиск DOI по библиографическим данным"""
        clean_ref = re.sub(r'\s*(https?://doi\.org/|doi:|DOI:)\s*[^\s,;]+', '', reference, flags=re.IGNORECASE)
        clean_ref = clean_ref.strip()
        
        if len(clean_ref) < 30:
            return None
        
        try:
            query = self.works.query(bibliographic=clean_ref).sort('relevance').order('desc')
            for result in query:
                if 'DOI' in result:
                    return result['DOI']
        except Exception as e:
            logger.error(f"Bibliographic search error for '{clean_ref}': {e}")
        
        return None
    
    def _find_openalex_doi(self, reference: str) -> Optional[str]:
        """Поиск DOI через OpenAlex API"""
        # Заглушка для будущей реализации OpenAlex
        # OpenAlex предоставляет бесплатный API с хорошими лимитами
        return None

    def extract_metadata_with_cache(self, doi: str) -> Optional[Dict]:
        """Извлечение метаданных с использованием кэша"""
        # Проверка кэша
        cached_metadata = self.cache.get(doi)
        if cached_metadata:
            logger.info(f"Cache hit for DOI: {doi}")
            return cached_metadata
        
        # Извлечение из API
        logger.info(f"Cache miss for DOI: {doi}, fetching from API")
        metadata = self._extract_metadata_from_api(doi)
        
        if metadata:
            self.cache.set(doi, metadata)
        
        return metadata
    
    def _extract_metadata_from_api(self, doi: str) -> Optional[Dict]:
        """Извлечение метаданных из Crossref API"""
        try:
            result = self.works.doi(doi)
            if not result:
                return None
            
            authors = result.get('author', [])
            author_list = []
            for author in authors:
                given_name = author.get('given', '')
                family_name = self._normalize_name(author.get('family', ''))
                author_list.append({
                    'given': given_name,
                    'family': family_name
                })
            
            title = ''
            if 'title' in result and result['title']:
                title = self._clean_text(result['title'][0])
                title = re.sub(r'</?sub>|</?i>|</?SUB>|</?I>', '', title, flags=re.IGNORECASE)
            
            journal = ''
            if 'container-title' in result and result['container-title']:
                journal = self._clean_text(result['container-title'][0])
            
            year = None
            if 'published' in result and 'date-parts' in result['published']:
                date_parts = result['published']['date-parts']
                if date_parts and date_parts[0]:
                    year = date_parts[0][0]
            
            volume = result.get('volume', '')
            issue = result.get('issue', '')
            pages = result.get('page', '')
            article_number = result.get('article-number', '')
            
            metadata = {
                'authors': author_list,
                'title': title,
                'journal': journal,
                'year': year,
                'volume': volume,
                'issue': issue,
                'pages': pages,
                'article_number': article_number,
                'doi': doi,
                'original_doi': doi
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting metadata for DOI {doi}: {e}")
            return None
    
    def _normalize_name(self, name: str) -> str:
        """Нормализует имя автора"""
        if not name:
            return ''
        
        if '-' in name or "'" in name or '’' in name:
            parts = re.split(r'([-\'’])', name)
            normalized_parts = []
            
            for i, part in enumerate(parts):
                if part in ['-', "'", '’']:
                    normalized_parts.append(part)
                else:
                    if part:
                        normalized_parts.append(part[0].upper() + part[1:].lower() if len(part) > 1 else part.upper())
            
            return ''.join(normalized_parts)
        else:
            if len(name) > 1:
                return name[0].upper() + name[1:].lower()
            else:
                return name.upper()
    
    def _clean_text(self, text: str) -> str:
        """Очищает текст от HTML тегов и entities"""
        if not text:
            return ""
        
        text = re.sub(r'<[^>]+>', '', text)
        text = html.unescape(text)
        text = re.sub(r'&[^;]+;', '', text)
        return text.strip()

# Основные функции обработки
class ReferenceProcessor:
    """Основной процессор для обработки ссылок"""
    
    def __init__(self):
        self.doi_processor = DOIProcessor()
        self.progress_manager = ProgressManager()
        self.validator = StyleValidator()
    
    def process_references(self, references: List[str], style_config: Dict, 
                         progress_container, status_container) -> Tuple[List, io.BytesIO, int, int, Dict]:
        """Обработка списка ссылок с отображением прогресса"""
        # Валидация
        is_valid, validation_messages = self.validator.validate_references_count(references)
        for msg in validation_messages:
            if "error" in msg.lower():
                st.error(msg)
            else:
                st.warning(msg)
        
        if not is_valid:
            return [], io.BytesIO(), 0, 0, {}
        
        doi_list = []
        formatted_refs = []
        doi_found_count = 0
        doi_not_found_count = 0
        
        # Сбор DOI для пакетной обработки
        valid_dois = []
        reference_doi_map = {}
        
        for i, ref in enumerate(references):
            if self.doi_processor._is_section_header(ref):
                doi_list.append(f"{ref} [SECTION HEADER - SKIPPED]")
                formatted_refs.append((ref, False, None))
                continue
                
            doi = self.doi_processor.find_doi_enhanced(ref)
            if doi:
                valid_dois.append(doi)
                reference_doi_map[i] = doi
                doi_list.append(doi)
            else:
                error_msg = f"{ref}\nПроверьте источник и добавьте DOI вручную." if st.session_state.current_language == 'ru' else f"{ref}\nPlease check this source and insert the DOI manually."
                doi_list.append(error_msg)
                formatted_refs.append((error_msg, True, None))
                doi_not_found_count += 1
        
        # Пакетная обработка DOI
        if valid_dois:
            self._process_doi_batch(valid_dois, reference_doi_map, references, 
                                  formatted_refs, doi_list, style_config,
                                  progress_container, status_container)
        
        # Подсчет статистики
        doi_found_count = len([ref for ref in formatted_refs if not ref[1] and ref[2]])
        
        # Поиск дубликатов
        duplicates_info = self._find_duplicates(formatted_refs)
        
        # Создание TXT файла
        txt_buffer = self._create_txt_file(doi_list)
        
        return formatted_refs, txt_buffer, doi_found_count, doi_not_found_count, duplicates_info
    
    def _process_doi_batch(self, valid_dois, reference_doi_map, references, 
                          formatted_refs, doi_list, style_config,
                          progress_container, status_container):
        """Пакетная обработка DOI"""
        status_container.info(get_text('batch_processing'))
        
        # Настройка прогресса
        self.progress_manager.start_processing(len(valid_dois))
        
        # Создаем прогресс-бар, который всегда будет виден
        progress_bar = progress_container.progress(0)
        status_display = status_container.empty()
        
        # Первая попытка обработки
        metadata_results = self._extract_metadata_batch(valid_dois, progress_bar, status_display)
        
        # Обработка результатов
        doi_to_metadata = dict(zip(valid_dois, metadata_results))
        
        for i, ref in enumerate(references):
            if i in reference_doi_map:
                doi = reference_doi_map[i]
                metadata = doi_to_metadata.get(doi)
                
                if metadata:
                    formatted_ref, is_error = self._format_reference(metadata, style_config)
                    formatted_refs.append((formatted_ref, is_error, metadata))
                else:
                    error_msg = self._create_error_message(ref, st.session_state.current_language)
                    doi_list[doi_list.index(doi)] = error_msg
                    formatted_refs.append((error_msg, True, None))
        
        # Обновление прогресса
        self._update_progress_display(progress_bar, status_display, len(valid_dois), len(valid_dois), 0)
    
    def _extract_metadata_batch(self, doi_list, progress_bar, status_display) -> List:
        """Пакетное извлечение метаданных"""
        results = [None] * len(doi_list)
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=Config.CROSSREF_WORKERS) as executor:
            future_to_index = {
                executor.submit(self.doi_processor.extract_metadata_with_cache, doi): i 
                for i, doi in enumerate(doi_list)
            }
            
            completed = 0
            for future in concurrent.futures.as_completed(future_to_index):
                index = future_to_index[future]
                try:
                    result = future.result(timeout=Config.REQUEST_TIMEOUT)
                    results[index] = result
                except Exception as e:
                    logger.error(f"Error processing DOI at index {index}: {e}")
                    results[index] = None
                
                completed += 1
                self._update_progress_display(progress_bar, status_display, completed, len(doi_list), 0)
        
        # Повторная попытка для неудачных запросов
        failed_indices = [i for i, result in enumerate(results) if result is None]
        if failed_indices:
            logger.info(f"Retrying {len(failed_indices)} failed DOI requests")
            self._retry_failed_requests(failed_indices, doi_list, results, progress_bar, status_display)
        
        return results
    
    def _retry_failed_requests(self, failed_indices, doi_list, results, progress_bar, status_display):
        """Повторная попытка обработки неудачных запросов"""
        completed = len(doi_list) - len(failed_indices)
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=Config.CROSSREF_RETRY_WORKERS) as executor:
            retry_futures = {}
            for index in failed_indices:
                doi = doi_list[index]
                future = executor.submit(self.doi_processor.extract_metadata_with_cache, doi)
                retry_futures[future] = index
            
            for future in concurrent.futures.as_completed(retry_futures):
                index = retry_futures[future]
                try:
                    result = future.result(timeout=Config.REQUEST_TIMEOUT)
                    results[index] = result
                except Exception as e:
                    logger.error(f"Error in retry processing DOI at index {index}: {e}")
                    results[index] = None
                
                completed += 1
                self._update_progress_display(progress_bar, status_display, completed, len(doi_list), len(failed_indices))
    
    def _update_progress_display(self, progress_bar, status_display, completed, total, errors):
        """Обновление отображения прогресса"""
        progress_info = self.progress_manager.get_progress_info()
        progress_ratio = completed / total if total > 0 else 0
        progress_color = self.progress_manager.get_progress_color(progress_ratio)
        
        # Обновляем прогресс-бар
        progress_bar.progress(progress_ratio)
        
        # Обновление стиля прогресс-бара с цветом
        progress_bar.markdown(f"""
            <style>
                .stProgress > div > div > div > div {{
                    background-color: {progress_color};
                }}
            </style>
        """, unsafe_allow_html=True)
        
        # Обновляем текст статуса
        status_text = f"Processed: {completed}/{total} | Errors: {errors}"
        if progress_info['time_remaining']:
            mins_remaining = int(progress_info['time_remaining'] / 60)
            status_text += f" | ETA: {mins_remaining} min"
        
        status_display.text(status_text)
    
    def _format_reference(self, metadata: Dict, style_config: Dict) -> Tuple[Any, bool]:
        """Форматирование ссылки"""
        formatter = CitationFormatterFactory.create_formatter(style_config)
        return formatter.format_reference(metadata, False)
    
    def _find_duplicates(self, formatted_refs: List) -> Dict[int, int]:
        """Поиск дубликатов ссылок"""
        seen_hashes = {}
        duplicates_info = {}
        
        for i, (elements, is_error, metadata) in enumerate(formatted_refs):
            if is_error or not metadata:
                continue
                
            ref_hash = self._generate_reference_hash(metadata)
            if not ref_hash:
                continue
                
            if ref_hash in seen_hashes:
                duplicates_info[i] = seen_hashes[ref_hash]
            else:
                seen_hashes[ref_hash] = i
        
        return duplicates_info
    
    def _generate_reference_hash(self, metadata: Dict) -> Optional[str]:
        """Генерация хеша для идентификации дубликатов"""
        if not metadata:
            return None
        
        hash_string = ""
        
        if metadata.get('authors'):
            authors_hash = "|".join(sorted([author.get('family', '').lower() for author in metadata['authors']]))
            hash_string += authors_hash + "||"
        
        title = metadata.get('title', '')[:50].lower()
        hash_string += title + "||"
        
        hash_string += (metadata.get('journal', '') + "||").lower()
        hash_string += str(metadata.get('year', '')) + "||"
        hash_string += metadata.get('volume', '') + "||"
        hash_string += metadata.get('pages', '') + "||"
        hash_string += self._normalize_doi(metadata.get('doi', ''))
        
        return hashlib.md5(hash_string.encode('utf-8')).hexdigest()
    
    def _normalize_doi(self, doi: str) -> str:
        """Нормализация DOI"""
        if not doi:
            return ""
        return re.sub(r'^(https?://doi\.org/|doi:|DOI:)', '', doi, flags=re.IGNORECASE).lower().strip()
    
    def _create_error_message(self, ref: str, language: str) -> str:
        """Создание сообщения об ошибке"""
        if language == 'ru':
            return f"{ref}\nПроверьте источник и добавьте DOI вручную."
        else:
            return f"{ref}\nPlease check this source and insert the DOI manually."
    
    def _create_txt_file(self, doi_list: List[str]) -> io.BytesIO:
        """Создание TXT файла со списком DOI"""
        output_txt_buffer = io.StringIO()
        for doi in doi_list:
            output_txt_buffer.write(f"{doi}\n")
        output_txt_buffer.seek(0)
        return io.BytesIO(output_txt_buffer.getvalue().encode('utf-8'))

# UI компоненты
class UIComponents:
    """Компоненты пользовательского интерфейса"""
    
    def __init__(self):
        self.user_prefs = UserPreferencesManager()
    
    def render_header(self):
        """Рендер заголовка и контролов с выпадающим меню"""
        col_title, col_lang, col_theme, col_view, col_menu = st.columns([2, 1.5, 1.5, 1.2, 0.8])
    
        with col_title:
            st.title(get_text('header'))
    
        with col_lang:
            self._render_language_selector()
    
        with col_theme:
            self._render_theme_selector()
    
        with col_view:
            self._render_view_selector()
    
        with col_menu:
            # Выпадающее меню для дополнительных действий
            with st.popover("⚙️"):
                st.markdown("**Actions**")
                self._render_clear_button()
                st.markdown("---")
                self._render_back_button()
    
    def _render_language_selector(self):
        """Рендер селектора языка"""
        languages = [
            ('English', 'en'),
            ('Русский', 'ru'), 
            ('Deutsch', 'de'),
            ('Español', 'es'),
            ('Italiano', 'it'),
            ('日本語', 'ja'),
            ('中文', 'zh')
        ]
        
        current_lang = st.session_state.current_language
        current_lang_name = next((name for name, code in languages if code == current_lang), 'English')
        
        selected_language = st.selectbox(
            get_text('language'),
            languages,
            format_func=lambda x: x[0],
            index=next(i for i, (_, code) in enumerate(languages) if code == current_lang),
            key="language_selector"
        )
        
        if selected_language[1] != st.session_state.current_language:
            self._save_current_state()
            st.session_state.current_language = selected_language[1]
            self._save_user_preferences()
            st.rerun()
    
    def _render_theme_selector(self):
        """Рендер селектора темы"""
        themes = [
            (get_text('light_theme'), 'light'),
            (get_text('dark_theme'), 'dark')
        ]
        
        selected_theme = st.radio(
            get_text('theme_selector'),
            themes,
            format_func=lambda x: x[0],
            index=0 if st.session_state.current_theme == 'light' else 1,
            key="theme_selector",
            horizontal=True
        )
        
        if selected_theme[1] != st.session_state.current_theme:
            self._save_current_state()
            st.session_state.current_theme = selected_theme[1]
            self._save_user_preferences()
            st.rerun()
    
    def _render_view_selector(self):
        """Рендер переключателя вида"""
        mobile_view = st.session_state.mobile_view
        view_label = get_text('mobile_view') if mobile_view else get_text('desktop_view')
        
        if st.button(view_label, key="view_selector", use_container_width=True):
            self._save_current_state()
            st.session_state.mobile_view = not st.session_state.mobile_view
            self._save_user_preferences()
            st.rerun()
    
    def _render_clear_button(self):
        """Рендер кнопки Clear с иконкой"""
        if st.button("🗑️", help="Clear all settings", key="clear_button", use_container_width=True):
            self._clear_all_settings()

    def _render_back_button(self):
        """Рендер кнопки Back с иконкой"""
        if st.session_state.previous_states:
            if st.button("↩️", help="Back to previous state", key="back_button", use_container_width=True):
                self._restore_previous_state()
    
    def _save_current_state(self):
        """Сохранение текущего состояния для кнопки Back"""
        if 'previous_states' not in st.session_state:
            st.session_state.previous_states = []
        
        # Сохраняем только основные настройки для экономии памяти
        current_state = {
            'current_language': st.session_state.current_language,
            'current_theme': st.session_state.current_theme,
            'mobile_view': st.session_state.mobile_view,
            'num': st.session_state.num,
            'auth': st.session_state.auth,
            'sep': st.session_state.sep,
            'etal': st.session_state.etal,
            'doi': st.session_state.doi,
            'doilink': st.session_state.doilink,
            'page': st.session_state.page,
            'punct': st.session_state.punct,
            'journal_style': st.session_state.journal_style,
            'use_and_checkbox': st.session_state.use_and_checkbox,
            'use_ampersand_checkbox': st.session_state.use_ampersand_checkbox,
            'gost_style': st.session_state.gost_style,
            'acs_style': st.session_state.acs_style,
            'rsc_style': st.session_state.rsc_style,
            'cta_style': st.session_state.cta_style,
            'timestamp': time.time()
        }
        
        # Сохраняем элементы конфигурации
        for i in range(8):
            for prop in ['el', 'it', 'bd', 'pr', 'sp']:
                key = f"{prop}{i}"
                current_state[key] = st.session_state[key]
        
        # Добавляем в стек и ограничиваем размер
        st.session_state.previous_states.append(current_state)
        if len(st.session_state.previous_states) > st.session_state.max_undo_steps:
            st.session_state.previous_states.pop(0)
    
    def _clear_all_settings(self):
        """Очистка всех настроек"""
        self._save_current_state()
        
        # Сброс основных настроек
        st.session_state.num = "No numbering"
        st.session_state.auth = "AA Smith"
        st.session_state.sep = ", "
        st.session_state.etal = 0
        st.session_state.doi = "10.10/xxx"
        st.session_state.doilink = True
        st.session_state.page = "122–128"
        st.session_state.punct = ""
        st.session_state.journal_style = '{Full Journal Name}'
        st.session_state.use_and_checkbox = False
        st.session_state.use_ampersand_checkbox = False
        
        # Сброс стилей
        st.session_state.gost_style = False
        st.session_state.acs_style = False
        st.session_state.rsc_style = False
        st.session_state.cta_style = False
        
        # Сброс элементов конфигурации
        for i in range(8):
            st.session_state[f"el{i}"] = ""
            st.session_state[f"it{i}"] = False
            st.session_state[f"bd{i}"] = False
            st.session_state[f"pr{i}"] = False
            st.session_state[f"sp{i}"] = ". "
        
        # Сброс данных
        st.session_state.output_text_value = ""
        st.session_state.show_results = False
        st.session_state.download_data = {}
        
        st.rerun()
    
    def _restore_previous_state(self):
        """Восстановление предыдущего состояния"""
        if not st.session_state.previous_states:
            st.warning("No previous state to restore")
            return
        
        previous_state = st.session_state.previous_states.pop()
        
        # Восстанавливаем основные настройки
        for key, value in previous_state.items():
            if key in st.session_state and key != 'timestamp':
                st.session_state[key] = value
        
        st.rerun()
    
    def _save_user_preferences(self):
        """Сохранение пользовательских предпочтений"""
        ip = self.user_prefs.get_user_ip()
        preferences = {
            'language': st.session_state.current_language,
            'theme': st.session_state.current_theme,
            'mobile_view': st.session_state.mobile_view
        }
        self.user_prefs.save_preferences(ip, preferences)
    
    def load_user_preferences(self):
        """Загрузка пользовательских предпочтений"""
        if not st.session_state.user_prefs_loaded:
            ip = self.user_prefs.get_user_ip()
            prefs = self.user_prefs.get_preferences(ip)
            
            st.session_state.current_language = prefs['language']
            st.session_state.current_theme = prefs['theme'] 
            st.session_state.mobile_view = prefs['mobile_view']
            st.session_state.user_prefs_loaded = True
    
    def apply_theme_styles(self):
        """Применение стилей темы"""
        theme = Config.THEMES[st.session_state.current_theme]
        
        st.markdown(f"""
            <style>
            .block-container {{
                padding: 0.2rem;
                background-color: {theme['background']};
                color: {theme['text']};
                font-family: {theme['font']};
            }}
            .stSelectbox, .stTextInput, .stNumberInput, .stCheckbox, .stRadio, .stFileUploader, .stTextArea {{
                margin-bottom: 0.02rem;
                background-color: {theme['secondaryBackground']};
                border: 1px solid {theme['border']};
                border-radius: 0.25rem;
            }}
            .stTextArea {{ 
                height: 40px !important; 
                font-size: 0.7rem; 
                background-color: {theme['secondaryBackground']};
                color: {theme['text']};
            }}
            .stButton > button {{ 
                width: 100%; 
                padding: 0.05rem; 
                font-size: 0.7rem; 
                margin: 0.02rem; 
                background-color: {theme['primary']};
                color: white;
                border: none;
                border-radius: 0.25rem;
            }}
            h1, h2, h3 {{
                color: {theme['text']} !important;
            }}
            h1 {{ font-size: 1.0rem; margin-bottom: 0.05rem; }}
            h2 {{ font-size: 0.9rem; margin-bottom: 0.05rem; }}
            h3 {{ font-size: 0.8rem; margin-bottom: 0.02rem; }}
            label {{ 
                font-size: 0.65rem !important; 
                color: {theme['text']} !important;
            }}
            .stMarkdown {{ 
                font-size: 0.65rem; 
                color: {theme['text']};
            }}
            .stCheckbox > label {{ 
                font-size: 0.6rem; 
                color: {theme['text']};
            }}
            .stRadio > label {{ 
                font-size: 0.65rem; 
                color: {theme['text']};
            }}
            .stDownloadButton > button {{ 
                font-size: 0.7rem; 
                padding: 0.05rem; 
                margin: 0.02rem; 
                background-color: {theme['primary']};
                color: white;
                border: none;
                border-radius: 0.25rem;
            }}
            .element-row {{ margin: 0.01rem; padding: 0.01rem; }}
            .processing-header {{ font-size: 0.8rem; font-weight: bold; margin-bottom: 0.1rem; }}
            .processing-status {{ font-size: 0.7rem; margin-bottom: 0.05rem; }}
            .compact-row {{ margin-bottom: 0.1rem; }}
            .guide-text {{ font-size: 0.55rem !important; line-height: 1.1; margin-bottom: 0.1rem; }}
            .guide-title {{ font-size: 0.7rem !important; font-weight: bold; margin-bottom: 0.1rem; }}
            .guide-step {{ font-size: 0.55rem !important; line-height: 1.1; margin-bottom: 0.1rem; }}
            .guide-note {{ font-size: 0.55rem !important; font-style: italic; line-height: 1.1; margin-bottom: 0.1rem; margin-left: 0.5rem; }}
            .card {{
                background-color: {theme['cardBackground']};
                padding: 0.5rem;
                border-radius: 0.5rem;
                border: 1px solid {theme['border']};
                margin-bottom: 0.5rem;
            }}
            
            /* Мобильные стили */
            @media (max-width: 768px) {{
                .block-container {{ padding: 0.1rem; }}
                .stSelectbox, .stTextInput, .stNumberInput {{ 
                    font-size: 0.8rem !important;
                    margin-bottom: 0.1rem;
                }}
                .stButton > button {{
                    font-size: 0.8rem !important;
                    padding: 0.3rem !important;
                    margin: 0.1rem !important;
                }}
                .stCheckbox > label {{
                    font-size: 0.7rem !important;
                }}
                h1 {{ font-size: 1.1rem !important; }}
                h2 {{ font-size: 1.0rem !important; }}
                h3 {{ font-size: 0.9rem !important; }}
            }}
            
            /* Десктоп стили */
            @media (min-width: 769px) {{
                .mobile-only {{ display: none; }}
            }}
            
            /* Мобильные только */
            @media (max-width: 768px) {{
                .desktop-only {{ display: none; }}
            }}
            </style>
        """, unsafe_allow_html=True)

    def render_style_presets(self):
        """Рендер пресетов стилей"""
        col_preset, col_info = st.columns([3, 1])
        with col_preset:
            st.markdown(f"**{get_text('style_presets')}**")
        with col_info:
            st.markdown(f"<span title='{get_text('style_preset_tooltip')}'>ℹ️</span>", unsafe_allow_html=True)
        
        # Создаем контейнер для кнопок стилей
        if st.session_state.mobile_view:
            # Мобильный вид - вертикальное расположение
            if st.session_state.current_language == 'ru':
                # Для русского языка показываем кнопку ГОСТ
                if st.button(get_text('gost_button'), use_container_width=True, key="gost_button"):
                    self._apply_gost_style()
            # Для всех языков показываем остальные стили
            if st.button(get_text('acs_button'), use_container_width=True, key="acs_button"):
                self._apply_acs_style()
            if st.button(get_text('rsc_button'), use_container_width=True, key="rsc_button"):
                self._apply_rsc_style()
            if st.button(get_text('cta_button'), use_container_width=True, key="cta_button"):
                self._apply_cta_style()
        else:
            # Десктоп вид - горизонтальное расположение
            if st.session_state.current_language == 'ru':
                # Для русского языка показываем все 4 кнопки
                col_gost, col_acs, col_rsc, col_cta = st.columns(4)
                
                with col_gost:
                    if st.button(get_text('gost_button'), use_container_width=True, key="gost_button"):
                        self._apply_gost_style()
                
                with col_acs:
                    if st.button(get_text('acs_button'), use_container_width=True, key="acs_button"):
                        self._apply_acs_style()
                
                with col_rsc:
                    if st.button(get_text('rsc_button'), use_container_width=True, key="rsc_button"):
                        self._apply_rsc_style()
                
                with col_cta:
                    if st.button(get_text('cta_button'), use_container_width=True, key="cta_button"):
                        self._apply_cta_style()
            else:
                # Для других языков показываем только 3 кнопки (без ГОСТ)
                col_acs, col_rsc, col_cta = st.columns(3)
                
                with col_acs:
                    if st.button(get_text('acs_button'), use_container_width=True, key="acs_button"):
                        self._apply_acs_style()
                
                with col_rsc:
                    if st.button(get_text('rsc_button'), use_container_width=True, key="rsc_button"):
                        self._apply_rsc_style()
                
                with col_cta:
                    if st.button(get_text('cta_button'), use_container_width=True, key="cta_button"):
                        self._apply_cta_style()
    
    def _apply_gost_style(self):
        """Применение стиля ГОСТ (только для русского языка)"""
        def apply_gost_callback():
            self._save_current_state()
            st.session_state.num = "No numbering"
            st.session_state.auth = "Smith AA"  # Обновляем формат авторов
            st.session_state.sep = ", "
            st.session_state.etal = 0
            st.session_state.use_and_checkbox = False
            st.session_state.use_ampersand_checkbox = False
            st.session_state.doi = "https://dx.doi.org/10.10/xxx"
            st.session_state.doilink = True
            st.session_state.page = "122-128"  # Обновляем формат страниц
            st.session_state.punct = ""
            st.session_state.journal_style = "{Full Journal Name}"  # Используем полное название
            
            for i in range(8):
                st.session_state[f"el{i}"] = ""
                st.session_state[f"it{i}"] = False
                st.session_state[f"bd{i}"] = False
                st.session_state[f"pr{i}"] = False
                st.session_state[f"sp{i}"] = ". "
            
            st.session_state.gost_style = True
            st.session_state.acs_style = False
            st.session_state.rsc_style = False
            st.session_state.cta_style = False
            st.session_state.style_applied = True
        
        # Вызываем callback и перезагружаем страницу
        apply_gost_callback()
        st.rerun()
    
    def _apply_acs_style(self):
        """Применение стиля ACS"""
        def apply_acs_callback():
            self._save_current_state()
            st.session_state.num = "No numbering"
            st.session_state.auth = "Smith, A.A."
            st.session_state.sep = "; "
            st.session_state.etal = 0
            st.session_state.use_and_checkbox = False
            st.session_state.use_ampersand_checkbox = False
            st.session_state.doi = "10.10/xxx"
            st.session_state.doilink = True
            st.session_state.page = "122–128"
            st.session_state.punct = "."
            st.session_state.journal_style = "{J. Abbr.}"
            
            for i in range(8):
                st.session_state[f"el{i}"] = ""
                st.session_state[f"it{i}"] = False
                st.session_state[f"bd{i}"] = False
                st.session_state[f"pr{i}"] = False
                st.session_state[f"sp{i}"] = ". "
            
            st.session_state.gost_style = False
            st.session_state.acs_style = True
            st.session_state.rsc_style = False
            st.session_state.cta_style = False
            st.session_state.style_applied = True
        
        apply_acs_callback()
        st.rerun()
    
    def _apply_rsc_style(self):
        """Применение стиля RSC"""
        def apply_rsc_callback():
            self._save_current_state()
            st.session_state.num = "No numbering"
            st.session_state.auth = "A.A. Smith"
            st.session_state.sep = ", "
            st.session_state.etal = 0
            st.session_state.use_and_checkbox = True
            st.session_state.use_ampersand_checkbox = False
            st.session_state.doi = "10.10/xxx"
            st.session_state.doilink = True
            st.session_state.page = "122"
            st.session_state.punct = "."
            st.session_state.journal_style = "{J. Abbr.}"
            
            for i in range(8):
                st.session_state[f"el{i}"] = ""
                st.session_state[f"it{i}"] = False
                st.session_state[f"bd{i}"] = False
                st.session_state[f"pr{i}"] = False
                st.session_state[f"sp{i}"] = ". "
            
            st.session_state.gost_style = False
            st.session_state.acs_style = False
            st.session_state.rsc_style = True
            st.session_state.cta_style = False
            st.session_state.style_applied = True
        
        apply_rsc_callback()
        st.rerun()
    
    def _apply_cta_style(self):
        """Применение стиля CTA"""
        def apply_cta_callback():
            self._save_current_state()
            st.session_state.num = "No numbering"
            st.session_state.auth = "Smith AA"
            st.session_state.sep = ", "
            st.session_state.etal = 0
            st.session_state.use_and_checkbox = False
            st.session_state.use_ampersand_checkbox = False
            st.session_state.doi = "doi:10.10/xxx"
            st.session_state.doilink = True
            st.session_state.page = "122–8"
            st.session_state.punct = ""
            st.session_state.journal_style = "{J Abbr}"
            
            for i in range(8):
                st.session_state[f"el{i}"] = ""
                st.session_state[f"it{i}"] = False
                st.session_state[f"bd{i}"] = False
                st.session_state[f"pr{i}"] = False
                st.session_state[f"sp{i}"] = ". "
            
            st.session_state.gost_style = False
            st.session_state.acs_style = False
            st.session_state.rsc_style = False
            st.session_state.cta_style = True
            st.session_state.style_applied = True
        
        apply_cta_callback()
        st.rerun()
    
    def render_general_settings(self):
        """Рендер общих настроек"""
        if st.session_state.mobile_view:
            # Мобильный вид - вертикальное расположение
            numbering_style = st.selectbox(
                get_text('numbering_style'), 
                Config.NUMBERING_STYLES, 
                key="num", 
                index=Config.NUMBERING_STYLES.index(st.session_state.num)
            )
            
            author_format = st.selectbox(
                get_text('author_format'), 
                Config.AUTHOR_FORMATS, 
                key="auth", 
                index=Config.AUTHOR_FORMATS.index(st.session_state.auth)
            )
            
            col_sep_etal = st.columns(2)
            with col_sep_etal[0]:
                author_separator = st.selectbox(
                    get_text('author_separator'), 
                    [", ", "; "], 
                    key="sep", 
                    index=[", ", "; "].index(st.session_state.sep)
                )
            with col_sep_etal[1]:
                et_al_limit = st.number_input(
                    get_text('et_al_limit'), 
                    min_value=0, 
                    step=1, 
                    key="etal", 
                    value=st.session_state.etal
                )
            
            # Остальные настройки...
        else:
            # Десктоп вид - стандартное расположение
            numbering_style = st.selectbox(
                get_text('numbering_style'), 
                Config.NUMBERING_STYLES, 
                key="num", 
                index=Config.NUMBERING_STYLES.index(st.session_state.num)
            )
            
            # Настройки авторов
            col_authors = st.columns([1, 1, 1])
            with col_authors[0]:
                author_format = st.selectbox(
                    get_text('author_format'), 
                    Config.AUTHOR_FORMATS, 
                    key="auth", 
                    index=Config.AUTHOR_FORMATS.index(st.session_state.auth)
                )
            with col_authors[1]:
                author_separator = st.selectbox(
                    get_text('author_separator'), 
                    [", ", "; "], 
                    key="sep", 
                    index=[", ", "; "].index(st.session_state.sep)
                )
            with col_authors[2]:
                et_al_limit = st.number_input(
                    get_text('et_al_limit'), 
                    min_value=0, 
                    step=1, 
                    key="etal", 
                    value=st.session_state.etal
                )
        
        # Общие для обоих видов
        col_and_amp = st.columns(2)
        with col_and_amp[0]:
            use_and_checkbox = st.checkbox(
                get_text('use_and'), 
                key="use_and_checkbox", 
                value=st.session_state.use_and_checkbox,
                disabled=st.session_state.use_ampersand_checkbox
            )
        with col_and_amp[1]:
            use_ampersand_checkbox = st.checkbox(
                get_text('use_ampersand'), 
                key="use_ampersand_checkbox", 
                value=st.session_state.use_ampersand_checkbox,
                disabled=st.session_state.use_and_checkbox
            )
        
        # Стиль журнала
        journal_style = st.selectbox(
            get_text('journal_style'),
            Config.JOURNAL_STYLES,
            key="journal_style",
            index=Config.JOURNAL_STYLES.index(st.session_state.journal_style),
            format_func=lambda x: {
                "{Full Journal Name}": get_text('full_journal_name'),
                "{J. Abbr.}": get_text('journal_abbr_with_dots'),
                "{J Abbr}": get_text('journal_abbr_no_dots')
            }[x]
        )
        
        # Настройки страниц
        current_page = st.session_state.page
        page_index = 3
        if current_page in Config.PAGE_FORMATS:
            page_index = Config.PAGE_FORMATS.index(current_page)
        
        page_format = st.selectbox(
            get_text('page_format'), 
            Config.PAGE_FORMATS, 
            key="page", 
            index=page_index
        )
        
        # Настройки DOI
        if st.session_state.mobile_view:
            doi_format = st.selectbox(
                get_text('doi_format'), 
                Config.DOI_FORMATS, 
                key="doi", 
                index=Config.DOI_FORMATS.index(st.session_state.doi)
            )
            doi_hyperlink = st.checkbox(
                get_text('doi_hyperlink'), 
                key="doilink", 
                value=st.session_state.doilink
            )
        else:
            col_doi = st.columns([2, 1])
            with col_doi[0]:
                doi_format = st.selectbox(
                    get_text('doi_format'), 
                    Config.DOI_FORMATS, 
                    key="doi", 
                    index=Config.DOI_FORMATS.index(st.session_state.doi)
                )
            with col_doi[1]:
                doi_hyperlink = st.checkbox(
                    get_text('doi_hyperlink'), 
                    key="doilink", 
                    value=st.session_state.doilink
                )
        
        # Конечная пунктуация
        final_punctuation = st.selectbox(
            get_text('final_punctuation'), 
            ["", "."], 
            key="punct", 
            index=["", "."].index(st.session_state.punct)
        )
    
    def render_element_configuration(self):
        """Рендер конфигурации элементов"""
        element_configs = []
        used_elements = set()
        
        st.markdown(
            f"<small>{get_text('element')} | {get_text('italic')} | {get_text('bold')} | {get_text('parentheses')} | {get_text('separator')}</small>", 
            unsafe_allow_html=True
        )
        
        for i in range(8):
            if st.session_state.mobile_view:
                # Мобильный вид - вертикальное расположение
                element = st.selectbox(
                    f"Element {i+1}", 
                    Config.AVAILABLE_ELEMENTS, 
                    key=f"el{i}", 
                    index=Config.AVAILABLE_ELEMENTS.index(st.session_state[f"el{i}"]) if st.session_state[f"el{i}"] in Config.AVAILABLE_ELEMENTS else 0
                )
                
                col_mobile = st.columns(4)
                with col_mobile[0]:
                    italic = st.checkbox(
                        get_text('italic'), 
                        key=f"it{i}", 
                        value=st.session_state[f"it{i}"]
                    )
                with col_mobile[1]:
                    bold = st.checkbox(
                        get_text('bold'), 
                        key=f"bd{i}", 
                        value=st.session_state[f"bd{i}"]
                    )
                with col_mobile[2]:
                    parentheses = st.checkbox(
                        get_text('parentheses'), 
                        key=f"pr{i}", 
                        value=st.session_state[f"pr{i}"]
                    )
                with col_mobile[3]:
                    separator = st.text_input(
                        get_text('separator'), 
                        value=st.session_state[f"sp{i}"], 
                        key=f"sp{i}"
                    )
            else:
                # Десктоп вид - горизонтальное расположение
                cols = st.columns([2, 1, 1, 1, 2])
                
                with cols[0]:
                    element = st.selectbox(
                        "", 
                        Config.AVAILABLE_ELEMENTS, 
                        key=f"el{i}", 
                        label_visibility="collapsed",
                        index=Config.AVAILABLE_ELEMENTS.index(st.session_state[f"el{i}"]) if st.session_state[f"el{i}"] in Config.AVAILABLE_ELEMENTS else 0
                    )
                
                with cols[1]:
                    italic = st.checkbox(
                        "", 
                        key=f"it{i}", 
                        help=get_text('italic'), 
                        value=st.session_state[f"it{i}"]
                    )
                
                with cols[2]:
                    bold = st.checkbox(
                        "", 
                        key=f"bd{i}", 
                        help=get_text('bold'), 
                        value=st.session_state[f"bd{i}"]
                    )
                
                with cols[3]:
                    parentheses = st.checkbox(
                        "", 
                        key=f"pr{i}", 
                        help=get_text('parentheses'), 
                        value=st.session_state[f"pr{i}"]
                    )
                
                with cols[4]:
                    separator = st.text_input(
                        "", 
                        value=st.session_state[f"sp{i}"], 
                        key=f"sp{i}", 
                        label_visibility="collapsed"
                    )
            
            if element and element not in used_elements:
                element_configs.append((
                    element, 
                    {
                        'italic': italic, 
                        'bold': bold, 
                        'parentheses': parentheses, 
                        'separator': separator
                    }
                ))
                used_elements.add(element)
        
        return element_configs
    
    def render_style_preview(self, style_config: Dict):
        """Рендер предпросмотра стиля"""
        # Интерактивный предпросмотр
        current_time = time.time()
        if current_time - st.session_state.get('last_style_update', 0) > 1:
            st.session_state.last_style_update = current_time
            
            preview_metadata = self._get_preview_metadata(style_config)
            if preview_metadata:
                preview_ref, _ = format_reference(preview_metadata, style_config, for_preview=True)
                preview_with_numbering = self._add_numbering(preview_ref, style_config)
                
                # Форматирование HTML для предпросмотра
                preview_html = self._format_preview_html(preview_with_numbering, style_config)
                st.markdown(f"<small>{get_text('example')} {preview_html}</small>", unsafe_allow_html=True)
    
    def _get_preview_metadata(self, style_config: Dict) -> Optional[Dict]:
        """Получение метаданных для предпросмотра"""
        if style_config.get('gost_style', False):
            return {
                'authors': [{'given': 'John A.', 'family': 'Smith'}, {'given': 'Alice B.', 'family': 'Doe'}],
                'title': 'Article Title',
                'journal': 'Journal of the American Chemical Society',
                'year': 2020,
                'volume': '15',
                'issue': '3',
                'pages': '122-128',
                'article_number': '',
                'doi': '10.1000/xyz123'
            }
        elif style_config.get('acs_style', False):
            return {
                'authors': [{'given': 'John A.', 'family': 'Smith'}, {'given': 'Alice B.', 'family': 'Doe'}],
                'title': 'Article Title',
                'journal': 'Journal of the American Chemical Society',
                'year': 2020,
                'volume': '15',
                'issue': '3',
                'pages': '122-128',
                'article_number': '',
                'doi': '10.1000/xyz123'
            }
        elif style_config.get('rsc_style', False):
            return {
                'authors': [{'given': 'John A.', 'family': 'Smith'}, {'given': 'Alice B.', 'family': 'Doe'}],
                'title': 'Article Title',
                'journal': 'Chemical Communications',
                'year': 2020,
                'volume': '15',
                'issue': '3',
                'pages': '122-128',
                'article_number': '',
                'doi': '10.1000/xyz123'
            }
        elif style_config.get('cta_style', False):
            return {
                'authors': [
                    {'given': 'Fei', 'family': 'He'}, 
                    {'given': 'Feng', 'family': 'Ma'},
                    {'given': 'Juan', 'family': 'Li'},
                    {'given': 'Tao', 'family': 'Li'},
                    {'given': 'Guangshe', 'family': 'Li'}
                ],
                'title': 'Effect of calcination temperature on the structural properties and photocatalytic activities of solvothermal synthesized TiO2 hollow nanoparticles',
                'journal': 'Ceramics International',
                'year': 2014,
                'volume': '40',
                'issue': '5',
                'pages': '6441-6446',
                'article_number': '',
                'doi': '10.1016/j.ceramint.2013.11.094'
            }
        elif style_config.get('elements'):
            return {
                'authors': [{'given': 'John A.', 'family': 'Smith'}, {'given': 'Alice B.', 'family': 'Doe'}],
                'title': 'Article Title',
                'journal': 'Journal of the American Chemical Society',
                'year': 2020,
                'volume': '15',
                'issue': '3',
                'pages': '122-128',
                'article_number': 'e12345',
                'doi': '10.1000/xyz123'
            }
        else:
            return None
    
    def _add_numbering(self, preview_ref: str, style_config: Dict) -> str:
        """Добавление нумерации к предпросмотру"""
        numbering = style_config['numbering_style']
        if numbering == "No numbering":
            return preview_ref
        elif numbering == "1":
            return f"1 {preview_ref}"
        elif numbering == "1.":
            return f"1. {preview_ref}"
        elif numbering == "1)":
            return f"1) {preview_ref}"
        elif numbering == "(1)":
            return f"(1) {preview_ref}"
        elif numbering == "[1]":
            return f"[1] {preview_ref}"
        else:
            return f"1. {preview_ref}"
    
    def _format_preview_html(self, preview_text: str, style_config: Dict) -> str:
        """Форматирование HTML для предпросмотра"""
        preview_html = preview_text
        
        if style_config.get('acs_style', False):
            preview_html = preview_html.replace("J. Am. Chem. Soc.", "<i>J. Am. Chem. Soc.</i>")
            preview_html = preview_html.replace("2020", "<b>2020</b>")
            preview_html = preview_html.replace("15", "<i>15</i>")
        elif style_config.get('rsc_style', False):
            preview_html = preview_html.replace("Chem. Commun.", "<i>Chem. Commun.</i>")
            preview_html = preview_html.replace("15", "<b>15</b>")
        
        return preview_html

    def render_data_input(self):
        """Рендер ввода данных"""
        input_method = st.radio(
            get_text('input_method'), 
            ['DOCX', 'Text' if st.session_state.current_language == 'en' else 'Текст'], 
            horizontal=True, 
            key="input_method"
        )
        
        if input_method == 'DOCX':
            uploaded_file = st.file_uploader(
                get_text('select_docx'), 
                type=['docx'], 
                label_visibility="collapsed", 
                key="docx_uploader"
            )
            return uploaded_file
        else:
            references_input = st.text_area(
                get_text('references'), 
                placeholder=get_text('enter_references'), 
                height=40, 
                label_visibility="collapsed", 
                key="references_input"
            )
            return references_input
    
    def render_data_output(self):
        """Рендер вывода данных"""
        output_method = st.radio(
            get_text('output_method'), 
            ['DOCX', 'Text' if st.session_state.current_language == 'en' else 'Текст'], 
            horizontal=True, 
            key="output_method"
        )
        
        if output_method == 'Text' if st.session_state.current_language == 'en' else 'Текст':
            output_text_value = st.session_state.output_text_value if st.session_state.show_results else ""
            st.text_area(
                get_text('results'), 
                value=output_text_value, 
                height=40, 
                disabled=True, 
                label_visibility="collapsed", 
                key="output_text"
            )
        
        return output_method

    def render_guide(self):
        """Рендер руководства"""
        st.markdown(f"<div class='guide-title'>{get_text('short_guide_title')}</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='guide-step'>{get_text('step_1')}</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='guide-note'>{get_text('step_1_note')}</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='guide-step'>{get_text('step_2')}</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='guide-step'>{get_text('step_3')}</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='guide-step'>{get_text('step_4')}</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='guide-step'>{get_text('step_5')}</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='guide-note'>{get_text('step_5_note')}</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='guide-step'>{get_text('step_6')}</div>", unsafe_allow_html=True)

# Основной класс приложения
class CitationStyleApp:
    """Основной класс приложения"""
    
    def __init__(self):
        self.processor = ReferenceProcessor()
        self.validator = StyleValidator()
        self.ui = UIComponents()
        init_session_state()
    
    def run(self):
        """Запуск приложения"""
        st.set_page_config(layout="wide")
    
        # Загрузка пользовательских предпочтений
        self.ui.load_user_preferences()
    
        # Обработка импортированного стиля (если есть)
        self._handle_imported_style()
    
        # Применение стилей темы
        self.ui.apply_theme_styles()
        
        # Рендер заголовка и контролов
        self.ui.render_header()
        
        # Основной макет в зависимости от вида
        if st.session_state.mobile_view:
            self._render_mobile_layout()
        else:
            self._render_desktop_layout()
    
    def _handle_imported_style(self):
        """Обработка импортированного стиля"""
        # Проверяем, есть ли импортированный стиль и нужно ли его применить
        if (st.session_state.get('imported_style') and 
            st.session_state.get('apply_imported_style') and 
            not st.session_state.get('style_import_processed')):

            # Применяем стиль
            self._apply_imported_style(st.session_state.imported_style)
            
            # Сбрасываем флаги
            st.session_state.apply_imported_style = False
            st.session_state.imported_style = None
            st.session_state.style_import_processed = True
            
            # Перезагружаем страницу для применения изменений
            st.rerun()
    
    def _render_mobile_layout(self):
        """Рендер мобильного макета"""
        with st.container():
            # Общие настройки
            st.subheader(get_text('general_settings'))
            self.ui.render_style_presets()
            self.ui.render_general_settings()
            
            # Конфигурация элементов
            st.subheader(get_text('element_config'))
            element_configs = self.ui.render_element_configuration()
            
            # Предпросмотр
            st.subheader(get_text('style_preview'))
            style_config = self._get_style_config(element_configs)
            self.ui.render_style_preview(style_config)
            
            # Ввод/вывод данных
            st.subheader(get_text('data_input'))
            input_data = self.ui.render_data_input()
            
            st.subheader(get_text('data_output'))
            output_method = self.ui.render_data_output()
            
            # Кнопка обработки
            if st.button(get_text('process'), use_container_width=True, key="process_button"):
                self._process_data(input_data, style_config, output_method)
            
            # Кнопки скачивания
            self._render_download_buttons(output_method)
            
            # Управление стилями
            self._render_style_management(style_config)
            
            # Руководство
            st.markdown("---")
            self.ui.render_guide()
    
    def _render_desktop_layout(self):
        """Рендер десктоп макета"""
        col1, col2, col3 = st.columns([1, 1, 1])
        
        with col1:
            self._render_general_settings_column()
        
        with col2:
            self._render_element_config_column()
        
        with col3:
            self._render_preview_and_io_column()
    
    def _render_general_settings_column(self):
        """Рендер колонки с общими настройками"""
        st.subheader(get_text('general_settings'))
        self.ui.render_style_presets()
        self.ui.render_general_settings()
    
    def _render_element_config_column(self):
        """Рендер колонки с конфигурацией элементов"""
        st.subheader(get_text('element_config'))
        element_configs = self.ui.render_element_configuration()
        
        # Руководство
        st.markdown("---")
        self.ui.render_guide()
        
        return element_configs
    
    def _render_preview_and_io_column(self):
        """Рендер колонки с предпросмотром и вводом/выводом"""
        # Сбор конфигурации стиля
        style_config = self._get_style_config()
        
        # Предпросмотр
        st.subheader(get_text('style_preview'))
        self.ui.render_style_preview(style_config)
        
        # Ввод данных
        st.subheader(get_text('data_input'))
        input_data = self.ui.render_data_input()
        
        # Вывод данных
        st.subheader(get_text('data_output'))
        output_method = self.ui.render_data_output()
        
        # Кнопка обработки
        if st.button(get_text('process'), use_container_width=True, key="process_button"):
            self._process_data(input_data, style_config, output_method)
        
        # Кнопки скачивания
        self._render_download_buttons(output_method)
        
        # Управление стилями
        self._render_style_management(style_config)
    
    def _get_style_config(self, element_configs=None):
        """Получение конфигурации стиля"""
        if element_configs is None:
            element_configs = []
            used_elements = set()
            
            for i in range(8):
                element = st.session_state[f"el{i}"]
                if element and element not in used_elements:
                    element_configs.append((
                        element, 
                        {
                            'italic': st.session_state[f"it{i}"],
                            'bold': st.session_state[f"bd{i}"],
                            'parentheses': st.session_state[f"pr{i}"],
                            'separator': st.session_state[f"sp{i}"]
                        }
                    ))
                    used_elements.add(element)
        
        return {
            'author_format': st.session_state.auth,
            'author_separator': st.session_state.sep,
            'et_al_limit': st.session_state.etal if st.session_state.etal > 0 else None,
            'use_and_bool': st.session_state.use_and_checkbox,
            'use_ampersand_bool': st.session_state.use_ampersand_checkbox,
            'doi_format': st.session_state.doi,
            'doi_hyperlink': st.session_state.doilink,
            'page_format': st.session_state.page,
            'final_punctuation': st.session_state.punct,
            'numbering_style': st.session_state.num,
            'journal_style': st.session_state.journal_style,
            'elements': element_configs,
            'gost_style': st.session_state.get('gost_style', False),
            'acs_style': st.session_state.get('acs_style', False),
            'rsc_style': st.session_state.get('rsc_style', False),
            'cta_style': st.session_state.get('cta_style', False)
        }
    
    def _process_data(self, input_data, style_config, output_method):
        """Обработка данных"""
        # Валидация стиля
        is_valid, validation_messages = self.validator.validate_style_config(style_config)
        for msg in validation_messages:
            if "error" in msg.lower():
                st.error(msg)
                return
            else:
                st.warning(msg)
        
        if not is_valid:
            st.error(get_text('validation_error_no_elements'))
            return
        
        # Обработка в зависимости от типа ввода
        try:
            if isinstance(input_data, str):  # Текстовый ввод
                self._process_text_input(input_data, style_config, output_method)
            else:  # DOCX ввод
                self._process_docx_input(input_data, style_config, output_method)
        except Exception as e:
            logger.error(f"Processing error: {e}")
            st.error(f"Processing error: {str(e)}")
    
    def _process_text_input(self, references_input, style_config, output_method):
        """Обработка текстового ввода"""
        if not references_input.strip():
            st.error(get_text('enter_references_error'))
            return
        
        references = [ref.strip() for ref in references_input.split('\n') if ref.strip()]
        st.write(f"**{get_text('found_references_text').format(len(references))}**")
        
        # Создаем контейнеры для прогресса
        progress_container = st.empty()
        status_container = st.empty()
        
        with st.spinner(get_text('processing')):
            formatted_refs, txt_bytes, doi_found_count, doi_not_found_count, duplicates_info = (
                self.processor.process_references(references, style_config, progress_container, status_container)
            )
            
            statistics = generate_statistics(formatted_refs)
            output_doc_buffer = DocumentGenerator.generate_document(
                formatted_refs, statistics, style_config, duplicates_info
            )
            
            self._handle_output(formatted_refs, txt_bytes, output_doc_buffer, 
                              doi_found_count, doi_not_found_count, output_method)
    
    def _process_docx_input(self, uploaded_file, style_config, output_method):
        """Обработка DOCX ввода"""
        if not uploaded_file:
            st.error(get_text('upload_file'))
            return
        
        # Создаем контейнеры для прогресса
        progress_container = st.empty()
        status_container = st.empty()
        
        with st.spinner(get_text('processing')):
            doc = Document(uploaded_file)
            references = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            st.write(f"**{get_text('found_references').format(len(references))}**")
            
            formatted_refs, txt_bytes, output_doc_buffer, doi_found_count, doi_not_found_count, statistics = (
                self._process_docx_references(references, style_config, progress_container, status_container)
            )
            
            self._handle_output(formatted_refs, txt_bytes, output_doc_buffer,
                              doi_found_count, doi_not_found_count, output_method)
    
    def _process_docx_references(self, references, style_config, progress_container, status_container):
        """Обработка ссылок из DOCX"""
        formatted_refs, txt_bytes, doi_found_count, doi_not_found_count, duplicates_info = (
            self.processor.process_references(references, style_config, progress_container, status_container)
        )
        
        statistics = generate_statistics(formatted_refs)
        output_doc_buffer = DocumentGenerator.generate_document(
            formatted_refs, statistics, style_config, duplicates_info
        )
        
        return formatted_refs, txt_bytes, output_doc_buffer, doi_found_count, doi_not_found_count, statistics
    
    def _handle_output(self, formatted_refs, txt_bytes, output_doc_buffer, 
                      doi_found_count, doi_not_found_count, output_method):
        """Обработка вывода"""
        # Статистика
        st.write(f"**{get_text('statistics').format(doi_found_count, doi_not_found_count)}**")
        
        # Подготовка текстового вывода
        if output_method == 'Text' if st.session_state.current_language == 'en' else 'Текст':
            output_text_value = self._format_text_output(formatted_refs, st.session_state.num)
            st.session_state.output_text_value = output_text_value
            st.session_state.show_results = True
        else:
            st.session_state.output_text_value = ""
            st.session_state.show_results = False
        
        # Сохранение данных для скачивания
        st.session_state.download_data = {
            'txt_bytes': txt_bytes,
            'output_doc_buffer': output_doc_buffer
        }
        
        st.rerun()
    
    def _format_text_output(self, formatted_refs, numbering_style):
        """Форматирование текстового вывода"""
        output_text_value = ""
        for i, (elements, is_error, metadata) in enumerate(formatted_refs):
            prefix = self._get_numbering_prefix(i, numbering_style)
            
            if is_error:
                output_text_value += f"{prefix}{elements}\n"
            else:
                if isinstance(elements, str):
                    output_text_value += f"{prefix}{elements}\n"
                else:
                    ref_str = ""
                    for j, element_data in enumerate(elements):
                        if len(element_data) == 6:
                            value, _, _, separator, _, _ = element_data
                            ref_str += value
                            if separator and j < len(elements) - 1:
                                ref_str += separator
                        else:
                            ref_str += str(element_data)
                    
                    output_text_value += f"{prefix}{ref_str}\n"
        
        return output_text_value
    
    def _get_numbering_prefix(self, index, numbering_style):
        """Получение префикса нумерации"""
        if numbering_style == "No numbering":
            return ""
        elif numbering_style == "1":
            return f"{index + 1} "
        elif numbering_style == "1.":
            return f"{index + 1}. "
        elif numbering_style == "1)":
            return f"{index + 1}) "
        elif numbering_style == "(1)":
            return f"({index + 1}) "
        elif numbering_style == "[1]":
            return f"[{index + 1}] "
        else:
            return f"{index + 1}. "
    
    def _render_download_buttons(self, output_method):
        """Рендер кнопок скачивания"""
        if st.session_state.download_data:
            if st.session_state.mobile_view:
                # Мобильный вид - вертикальное расположение
                st.download_button(
                    label=get_text('doi_txt'),
                    data=st.session_state.download_data['txt_bytes'],
                    file_name='doi_list.txt',
                    mime='text/plain',
                    key="doi_download",
                    use_container_width=True
                )
                
                if output_method == 'DOCX' and st.session_state.download_data.get('output_doc_buffer'):
                    st.download_button(
                        label=get_text('references_docx'),
                        data=st.session_state.download_data['output_doc_buffer'],
                        file_name='Reformatted references.docx',
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        key="docx_download",
                        use_container_width=True
                    )
            else:
                # Десктоп вид - горизонтальное расположение
                col_download = st.columns(2)
                with col_download[0]:
                    st.download_button(
                        label=get_text('doi_txt'),
                        data=st.session_state.download_data['txt_bytes'],
                        file_name='doi_list.txt',
                        mime='text/plain',
                        key="doi_download",
                        use_container_width=True
                    )
                
                with col_download[1]:
                    if output_method == 'DOCX' and st.session_state.download_data.get('output_doc_buffer'):
                        st.download_button(
                            label=get_text('references_docx'),
                            data=st.session_state.download_data['output_doc_buffer'],
                            file_name='Reformatted references.docx',
                            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                            key="docx_download",
                            use_container_width=True
                        )
    
    def _render_style_management(self, style_config):
        """Рендер управления стилями"""
        st.subheader("💾 Style Management")
        
        # Экспорт стиля
        col_export = st.columns([2, 1])
        with col_export[0]:
            export_file_name = st.text_input(
                get_text('export_file_name'), 
                value="my_citation_style", 
                placeholder="Enter file name", 
                key="export_name"
            )
        
        with col_export[1]:
            export_data = self._export_style(style_config, export_file_name)
            if export_data:
                st.download_button(
                    label=get_text('export_style'),
                    data=export_data,
                    file_name=f"{export_file_name}.json",
                    mime="application/json",
                    use_container_width=True,
                    key="export_button"
                )
        
        # Импорт стиля
        imported_file = st.file_uploader(
            get_text('import_file'), 
            type=['json'], 
            label_visibility="collapsed", 
            key="style_importer"
        )
        
        # Обработка импортированного файла
        if imported_file is not None:
            current_file_hash = hashlib.md5(imported_file.getvalue()).hexdigest()
            
            # Проверяем, что файл еще не обрабатывался в этой сессии
            if (st.session_state.last_imported_file_hash != current_file_hash or 
                not st.session_state.style_import_processed):
                
                imported_style = self._import_style(imported_file)
                if imported_style:
                    # Сохраняем хеш файла и устанавливаем флаги для применения
                    st.session_state.last_imported_file_hash = current_file_hash
                    st.session_state.imported_style = imported_style
                    st.session_state.apply_imported_style = True
                    st.session_state.style_import_processed = False
                    
                    st.success(get_text('import_success'))
                    st.rerun()
    
    def _render_cache_management(self):
        """Рендер управления кэшем"""
        with st.expander("Cache Management"):
            col_cache = st.columns(2)
            with col_cache[0]:
                if st.button("Initialize Cache", use_container_width=True):
                    doi_cache.clear_old_entries()
                    st.success(get_text('cache_initialized'))
            with col_cache[1]:
                if st.button("Clear Cache", use_container_width=True):
                    try:
                        os.remove(Config.DB_PATH)
                        st.success(get_text('cache_cleared'))
                    except Exception as e:
                        st.error(f"Error clearing cache: {e}")
    
    def _export_style(self, style_config, file_name):
        """Экспорт стиля"""
        try:
            export_data = {
                'version': '1.0',
                'export_date': str(datetime.now()),
                'style_config': style_config
            }
            json_data = json.dumps(export_data, indent=2, ensure_ascii=False)
            return json_data.encode('utf-8')
        except Exception as e:
            st.error(f"Export error: {str(e)}")
            return None
    
    def _import_style(self, uploaded_file):
        """Импорт стиля"""
        try:
            # Сохраняем позицию файла для возможности повторного чтения
            uploaded_file.seek(0)
            content = uploaded_file.read().decode('utf-8')
            import_data = json.loads(content)
        
            # Поддержка разных форматов файлов
            if 'style_config' in import_data:
                return import_data['style_config']
            elif 'version' in import_data:
                return import_data.get('style_config', import_data)
            else:
                # Предполагаем, что это прямой style_config
                return import_data
            
        except Exception as e:
            st.error(f"{get_text('import_error')}: {str(e)}")
            return None

    def _apply_imported_style(self, imported_style):
        """Применение импортированного стиля к состоянию сессии"""
        if not imported_style:
            return
    
        # Используем безопасный подход с callback для обновления состояния
        def apply_style_callback():
            # Общие настройки
            if 'numbering_style' in imported_style:
                st.session_state.num = imported_style['numbering_style']
            if 'author_format' in imported_style:
                st.session_state.auth = imported_style['author_format']
            if 'author_separator' in imported_style:
                st.session_state.sep = imported_style['author_separator']
            if 'et_al_limit' in imported_style:
                st.session_state.etal = imported_style['et_al_limit'] or 0
            if 'use_and_bool' in imported_style:
                st.session_state.use_and_checkbox = imported_style['use_and_bool']
            if 'use_ampersand_bool' in imported_style:
                st.session_state.use_ampersand_checkbox = imported_style['use_ampersand_bool']
            if 'doi_format' in imported_style:
                st.session_state.doi = imported_style['doi_format']
            if 'doi_hyperlink' in imported_style:
                st.session_state.doilink = imported_style['doi_hyperlink']
            if 'page_format' in imported_style:
                st.session_state.page = imported_style['page_format']
            if 'final_punctuation' in imported_style:
                st.session_state.punct = imported_style['final_punctuation']
            if 'journal_style' in imported_style:
                st.session_state.journal_style = imported_style['journal_style']
        
            # Сброс пресетов стилей
            st.session_state.gost_style = imported_style.get('gost_style', False)
            st.session_state.acs_style = imported_style.get('acs_style', False)
            st.session_state.rsc_style = imported_style.get('rsc_style', False)
            st.session_state.cta_style = imported_style.get('cta_style', False)
        
            # Очистка элементов
            for i in range(8):
                st.session_state[f"el{i}"] = ""
                st.session_state[f"it{i}"] = False
                st.session_state[f"bd{i}"] = False
                st.session_state[f"pr{i}"] = False
                st.session_state[f"sp{i}"] = ". "
        
            # Применение элементов из импортированного стиля
            elements = imported_style.get('elements', [])
            for i, (element, config) in enumerate(elements):
                if i < 8:  # Ограничиваем 8 элементами
                    st.session_state[f"el{i}"] = element
                    st.session_state[f"it{i}"] = config.get('italic', False)
                    st.session_state[f"bd{i}"] = config.get('bold', False)
                    st.session_state[f"pr{i}"] = config.get('parentheses', False)
                    st.session_state[f"sp{i}"] = config.get('separator', ". ")
        
            st.session_state.style_applied = True
            st.session_state.style_import_processed = True
        
        # Вызываем callback
        apply_style_callback()

# Вспомогательные функции
def clean_text(text):
    return DOIProcessor()._clean_text(text)

def normalize_name(name):
    return DOIProcessor()._normalize_name(name)

def is_section_header(text):
    return DOIProcessor()._is_section_header(text)

def find_doi(reference):
    return DOIProcessor().find_doi_enhanced(reference)

def normalize_doi(doi):
    processor = ReferenceProcessor()
    return processor._normalize_doi(doi)

def generate_reference_hash(metadata):
    processor = ReferenceProcessor()
    return processor._generate_reference_hash(metadata)

def extract_metadata_batch(doi_list, progress_callback=None):
    processor = ReferenceProcessor()
    return [processor.doi_processor.extract_metadata_with_cache(doi) for doi in doi_list]

def extract_metadata_sync(doi):
    processor = ReferenceProcessor()
    return processor.doi_processor.extract_metadata_with_cache(doi)

def format_reference(metadata, style_config, for_preview=False):
    formatter = CitationFormatterFactory.create_formatter(style_config)
    return formatter.format_reference(metadata, for_preview)

def find_duplicate_references(formatted_refs):
    processor = ReferenceProcessor()
    return processor._find_duplicates(formatted_refs)

def generate_statistics(formatted_refs):
    journals = []
    years = []
    authors = []
    
    current_year = datetime.now().year
    
    for _, _, metadata in formatted_refs:
        if not metadata:
            continue
            
        if metadata.get('journal'):
            journals.append(metadata['journal'])
        
        if metadata.get('year'):
            years.append(metadata['year'])
        
        if metadata.get('authors'):
            for author in metadata['authors']:
                given = author.get('given', '')
                family = author.get('family', '')
                if family:
                    first_initial = given[0] if given else ''
                    author_formatted = f"{family} {first_initial}." if first_initial else family
                    authors.append(author_formatted)
    
    unique_dois = set()
    for _, _, metadata in formatted_refs:
        if metadata and metadata.get('doi'):
            unique_dois.add(metadata['doi'])
    
    total_unique_dois = len(unique_dois)
    
    journal_counter = Counter(journals)
    journal_stats = []
    for journal, count in journal_counter.most_common(20):
        percentage = (count / total_unique_dois) * 100 if total_unique_dois > 0 else 0
        journal_stats.append({
            'journal': journal,
            'count': count,
            'percentage': round(percentage, 2)
        })
    
    year_counter = Counter(years)
    year_stats = []
    for year in range(current_year, 2009, -1):
        if year in year_counter:
            count = year_counter[year]
            percentage = (count / total_unique_dois) * 100 if total_unique_dois > 0 else 0
            year_stats.append({
                'year': year,
                'count': count,
                'percentage': round(percentage, 2)
            })
    
    recent_years = [current_year - i for i in range(4)]
    recent_count = sum(year_counter.get(year, 0) for year in recent_years)
    recent_percentage = (recent_count / total_unique_dois) * 100 if total_unique_dois > 0 else 0
    needs_more_recent_references = recent_percentage < 20
    
    author_counter = Counter(authors)
    author_stats = []
    for author, count in author_counter.most_common(20):
        percentage = (count / total_unique_dois) * 100 if total_unique_dois > 0 else 0
        author_stats.append({
            'author': author,
            'count': count,
            'percentage': round(percentage, 2)
        })
    
    has_frequent_author = any(stats['percentage'] > 30 for stats in author_stats)
    
    return {
        'journal_stats': journal_stats,
        'year_stats': year_stats,
        'author_stats': author_stats,
        'total_unique_dois': total_unique_dois,
        'needs_more_recent_references': needs_more_recent_references,
        'has_frequent_author': has_frequent_author
    }

def process_references_with_progress(references, style_config, progress_container, status_container):
    processor = ReferenceProcessor()
    return processor.process_references(references, style_config, progress_container, status_container)

def process_docx(input_file, style_config, progress_container, status_container):
    processor = ReferenceProcessor()
    doc = Document(input_file)
    references = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    return processor.process_references(references, style_config, progress_container, status_container)

def export_style(style_config, file_name):
    app = CitationStyleApp()
    return app._export_style(style_config, file_name)

def import_style(uploaded_file):
    app = CitationStyleApp()
    return app._import_style(uploaded_file)

def apply_imported_style(imported_style):
    """Функция для применения импортированного стиля (для обратной совместимости)"""
    app = CitationStyleApp()
    app._apply_imported_style(imported_style)

def main():
    """Основная функция"""
    app = CitationStyleApp()
    app.run()

if __name__ == "__main__":
    main()

